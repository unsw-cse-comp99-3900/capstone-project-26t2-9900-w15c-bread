from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "milestone2-text-diff-deep-qa-bilingual-2026-07-22.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GREY = "F4F6F9"
MID_GREY = "667085"
TEXT = "1F2937"
YELLOW = "FFD500"  # Named UNSW milestone accent override.
RED = "B42318"
GREEN = "067647"
USABLE_DXA = 9360  # 6.5 inches: US Letter width minus 1-inch margins.


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Iterable[int], indent=120) -> None:
    widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header_text(section, text: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.runs[0]
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MID_GREY)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), YELLOW)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)


def add_bottom_rule(paragraph, color="D0D5DD") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Title", "Subtitle"):
        styles[style_name].font.name = "Calibri"

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for name, color, bold in (
        ("Label EN", BLUE, True),
        ("Label CN", GREEN, True),
        ("Evidence", MID_GREY, True),
        ("Pitfall", RED, True),
    ):
        if name not in styles:
            st = styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        else:
            st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(10)
        st.font.bold = bold
        st.font.color.rgb = RGBColor.from_string(color)

    list_style = styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_callout(doc: Document, title: str, body: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_labelled_paragraph(doc: Document, label: str, text: str, label_style: str, *, final=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5 if not final else 10)
    p.paragraph_format.keep_with_next = not final
    r = p.add_run(label)
    r.style = doc.styles[label_style]
    p.add_run(text)
    if final:
        add_bottom_rule(p)


def add_question(doc: Document, number: int, item: dict) -> None:
    # These longer questions otherwise start at the bottom of a page in Word's
    # paginator. Explicit breaks keep every bilingual answer and its evidence
    # together, which matters more here than maximising page density.
    if number in {2, 20, 46}:
        doc.add_page_break()
    h = doc.add_paragraph(style="Heading 3")
    h.paragraph_format.keep_with_next = True
    h.add_run(f"Q{number}. {item['q_en']}")
    add_labelled_paragraph(doc, "中文问题：", item["q_cn"], "Label CN")
    add_labelled_paragraph(doc, "Answer (EN): ", item["a_en"], "Label EN")
    add_labelled_paragraph(doc, "回答（中文）：", item["a_cn"], "Label CN")
    add_labelled_paragraph(doc, "Code evidence / 代码证据：", item["evidence"], "Evidence")
    add_labelled_paragraph(doc, "Do not say / 易错点：", item["pitfall"], "Pitfall", final=True)


SECTIONS = [
    (
        "1. Milestone 2 narrative / Milestone 2 项目叙事",
        [
            {
                "q_en": "What is the core Milestone 2 deliverable?",
                "q_cn": "Milestone 2 的核心交付成果是什么？",
                "a_en": "Milestone 2 turns Dynamic History from a version-viewing idea into an interactive, block-aware text-diff and selective-recovery workflow for Confluence. A user can compare a historical version with the current page, inspect meaningful rich-text changes, choose which changes to keep or restore, preview the composed result, and publish it safely.",
                "a_cn": "Milestone 2 把 Dynamic History 从“查看历史版本”的概念推进为 Confluence 中可交互、理解区块结构的文本差异与选择性恢复流程。用户可以比较历史版本和当前页面，查看有语义的富文本变化，选择保留或恢复哪些内容，预览组合结果，并安全发布。",
                "evidence": "UI flow in App.js and TextDiffPage.js; core composition in buildRichTextDiffHtml (utils.js:6026–6115); publish path rechecks the live page before update.",
                "pitfall": "Do not describe it as a complete page rollback tool. / 不要把它说成只能整页回滚的工具。",
            },
            {
                "q_en": "How did the project evolve from Sprint 1 to Sprint 2?",
                "q_cn": "项目从 Sprint 1 到 Sprint 2 是如何演进的？",
                "a_en": "Sprint 1 established the end-to-end path: retrieve versions, compare old and current content, render inline differences, and support basic recovery. Sprint 2 hardened the difficult cases—lists, blank lines, tasks, decisions, tables, layouts, code, stable selections, side-by-side rendering, storage reconstruction, and safer write-back. The change was mainly from a prototype diff to a structured recovery system.",
                "a_cn": "Sprint 1 建立了端到端主路径：获取版本、比较旧内容与当前内容、展示行内差异并支持基础恢复。Sprint 2 重点加固困难场景，包括列表、空行、任务、决策、表格、布局、代码、稳定选择、并排视图、Storage 重建和更安全的写回。核心演进是从“差异原型”变为“结构化恢复系统”。",
                "evidence": "The current test suite contains dedicated suites for tables, layouts, tasks/decisions, blank lines, sanitisation, inline diff, side-by-side view, and TextDiffPage recovery.",
                "pitfall": "Do not claim every rich-text feature was newly invented in Sprint 2 unless the commit history proves it. / 不要在没有提交历史证据时声称所有功能都由 Sprint 2 首次实现。",
            },
            {
                "q_en": "What user problem does text diff solve?",
                "q_cn": "Text diff 解决了什么用户问题？",
                "a_en": "Confluence’s normal version history is useful for inspection, but restoring usually operates at page-version level. Our text diff reduces the cost of understanding long-page changes and allows users to recover only the parts they need, which is valuable when useful new edits and accidentally removed old content coexist.",
                "a_cn": "Confluence 原生版本历史适合查看，但恢复通常以整个页面版本为单位。我们的 text diff 降低了理解长页面变化的成本，并允许用户只恢复需要的部分；当“有价值的新编辑”和“误删的旧内容”同时存在时，这一点尤其重要。",
                "evidence": "The decision model defaults each changed block to current and lets the user explicitly choose old, producing a mixed preview rather than replacing the entire page.",
                "pitfall": "Do not say Confluence has no history or restore function. The gap is selective, structure-aware recovery. / 不要说 Confluence 没有历史或恢复；差异点是选择性、结构感知的恢复。",
            },
            {
                "q_en": "Who is the primary user and what is the main scenario?",
                "q_cn": "主要用户是谁，核心使用场景是什么？",
                "a_en": "The primary user is a Confluence editor, project lead, or knowledge owner who needs to audit a page and recover specific lost or overwritten content without discarding later work. The main scenario is selecting an older version, comparing it against the live version, restoring chosen fragments, validating the preview, and publishing a new version.",
                "a_cn": "主要用户是需要审查页面并恢复特定丢失或被覆盖内容的 Confluence 编辑者、项目负责人或知识库维护者，同时又不能丢弃后续工作。核心场景是选择旧版本，与当前线上版本比较，恢复指定片段，检查预览，然后发布为新版本。",
                "evidence": "The UI’s version selector, diff view, bulk actions, preview panel, and save/publish action form this workflow.",
                "pitfall": "Avoid framing the app as a developer-only source diff. / 不要把该应用描述成只给开发者使用的源码 diff。",
            },
            {
                "q_en": "How does the presentation’s research scope relate to the implementation?",
                "q_cn": "PPT 中的研究范围与实际实现是什么关系？",
                "a_en": "The presentation’s figures—such as considering more than 80 content types and prioritising 34—describe the design and research scope. The implementation focuses on high-value structures and provides generic preservation paths for unsupported content. Those research numbers should not be presented as a claim that all 34 types have identical, fully tested recovery semantics.",
                "a_cn": "PPT 中“考虑 80 多种内容类型、优先处理 34 种”等数字描述的是设计和调研范围。实际实现聚焦高价值结构，并为未专门支持的内容提供通用保留路径。不能把这些调研数字说成“34 种类型都拥有完全相同且经过完整测试的恢复语义”。",
                "evidence": "Specialised code exists for several structures, while raw-block fallbacks preserve other Confluence storage fragments.",
                "pitfall": "Do not convert research counts into runtime guarantees. / 不要把调研数量等同于运行时能力保证。",
            },
            {
                "q_en": "What would you define as success for Milestone 2?",
                "q_cn": "你会如何定义 Milestone 2 的成功？",
                "a_en": "Success means the core selective-recovery workflow works end to end, important rich-text structures remain understandable and recoverable, the app refuses unsafe writes, and automated tests cover the most failure-prone cases. It does not require perfect semantic interpretation of every Confluence macro.",
                "a_cn": "成功意味着选择性恢复主流程可以端到端运行，重要富文本结构能够被理解和恢复，应用会拒绝不安全写入，并且自动化测试覆盖最容易失败的场景。它不要求对每一种 Confluence 宏都做到完美语义解释。",
                "evidence": "Current local result: 11 test suites and 169 tests passing at commit 2e3e05718c37968df25642b68c2c82db39ebfde4.",
                "pitfall": "Do not use only the number of UI features as the success metric. / 不要只用 UI 功能数量衡量成功。",
            },
        ],
    ),
    (
        "2. Diff pipeline and algorithm / 差异管线与算法",
        [
            {
                "q_en": "What is the exact direction of the comparison?",
                "q_cn": "差异比较的准确方向是什么？",
                "a_en": "The selected historical version is the old input and the live page is the current input. Therefore, removed means present in history but absent from current, while added means absent from history but present in current. Recovery then decides whether a changed unit should contribute old or current storage content.",
                "a_cn": "用户选择的历史版本是 old 输入，线上页面是 current 输入。因此 removed 表示“历史中存在、当前已不存在”，added 表示“历史中不存在、当前新增”。恢复阶段再决定每个变化单元应使用旧版还是当前版的 Storage 内容。",
                "evidence": "TextDiffPage builds the comparison in historical → current order and passes both storage representations into the diff builder.",
                "pitfall": "Do not reverse added and removed during the demo. / 演示时不要把 added 和 removed 方向说反。",
            },
            {
                "q_en": "Can you explain the pipeline in one concise sequence?",
                "q_cn": "你能用一条简洁流程解释整个管线吗？",
                "a_en": "Fetch old and current versions, obtain rendered HTML plus raw Storage XHTML, sanitise the rendered preview, extract semantic blocks, create canonical signatures, align block sequences, compute finer differences inside matched structures, attach stable decision keys, render inline or side-by-side views, compose a preview from user decisions, then reconstruct and publish Storage XHTML.",
                "a_cn": "先获取旧版和当前版，同时取得渲染 HTML 与原始 Storage XHTML；清理用于展示的 HTML；提取语义区块；生成规范签名；对齐区块序列；在匹配结构内部计算更细差异；附加稳定决策键；渲染行内或并排视图；根据用户选择组合预览；最后重建并发布 Storage XHTML。",
                "evidence": "prepareConfluenceHtml → block extraction/signatures → buildRichTextDiffHtml → diffDisplay decisions → preview composition → recovery write-back.",
                "pitfall": "Do not collapse preview HTML and publishable Storage into one representation. / 不要把预览 HTML 和可发布 Storage 混为同一种表示。",
            },
            {
                "q_en": "Why is the diff block-aware instead of comparing one plain-text string?",
                "q_cn": "为什么 diff 要理解区块，而不是比较一整段纯文本？",
                "a_en": "Plain text discards hierarchy and Confluence semantics. Block-aware comparison keeps paragraphs, headings, list items, task items, decisions, tables, layouts, code, images, and blank lines meaningful enough for display and recovery. It also prevents a small change in one structure from causing unrelated parts of the page to appear changed.",
                "a_cn": "纯文本会丢失层级和 Confluence 语义。区块感知比较让段落、标题、列表项、任务、决策、表格、布局、代码、图片和空行在展示与恢复中仍具意义，也能避免某个结构中的小改动让无关页面区域看起来全部变化。",
                "evidence": "Semantic extraction is implemented around utils.js:3803–4192, with specialised table/layout/task branches.",
                "pitfall": "Do not claim it builds a complete formal AST for all Confluence content. / 不要声称它为所有 Confluence 内容建立了完整形式化 AST。",
            },
            {
                "q_en": "What is a canonical signature and why is it needed?",
                "q_cn": "什么是 canonical signature，为什么需要它？",
                "a_en": "A canonical signature is a normalised representation used to decide whether two blocks are semantically equivalent for alignment. It removes presentation noise and standardises relevant structure so that inconsequential HTML differences do not create false changes. The raw source is still retained separately for recovery.",
                "a_cn": "Canonical signature 是用于判断两个区块在对齐时是否语义等价的规范化表示。它去除展示噪声并标准化相关结构，避免无意义的 HTML 差异制造假变化；与此同时，原始内容仍单独保留用于恢复。",
                "evidence": "Canonicalisation and signature construction are concentrated in utils.js:3210–3569.",
                "pitfall": "Do not say canonicalisation overwrites the source content. / 不要说规范化会覆盖源内容。",
            },
            {
                "q_en": "Why is Longest Common Subsequence used?",
                "q_cn": "为什么使用最长公共子序列（LCS）？",
                "a_en": "LCS gives an understandable order-preserving alignment between old and current block sequences. Blocks present in both remain stable, while unmatched old and current blocks become removed and added candidates. That model works well for document history because reading order matters and the result is straightforward to convert into a diff stream.",
                "a_cn": "LCS 能在旧版和当前版区块序列之间建立保持顺序、容易解释的对齐。两边共有区块保持稳定，未匹配的旧区块和当前区块分别成为 removed 与 added 候选。文档历史强调阅读顺序，因此这种模型非常合适，也容易转换成差异流。",
                "evidence": "The page-level alignment has a 120,000-cell matrix cap; smaller inline helpers use their own caps.",
                "pitfall": "Do not say LCS detects moves as a first-class operation. Reordering can appear as remove plus add. / 不要说 LCS 原生识别移动；重排可能表现为删除加新增。",
            },
            {
                "q_en": "What are the formal page-level diff types?",
                "q_cn": "页面级 diff 的正式类型有哪些？",
                "a_en": "The main diff builder emits same, removed, and added. A visible modification is represented structurally as an old removed block followed by a related new added block. The display layer can pair those related entries and present one modified decision to the user.",
                "a_cn": "主 diff 构建器输出 same、removed 和 added。可见的“修改”在结构上表示为旧内容的 removed，后接相关新内容的 added；展示层可以把这两个相关条目配对，并向用户呈现一个 modified 决策。",
                "evidence": "buildRichTextDiffHtml emits the formal stream; diffDisplay.js performs related removed/added pairing for display decisions.",
                "pitfall": "Do not claim modified is a fourth page-level primitive emitted directly by the main builder. / 不要声称 modified 是主构建器直接输出的第四种页面级原语。",
            },
            {
                "q_en": "How are removed and added blocks paired as a modification?",
                "q_cn": "removed 和 added 区块如何被配对成一次修改？",
                "a_en": "The display layer evaluates whether neighbouring or related removed and added blocks are similar enough to represent one replacement. It uses normalised content and similarity signals, including character-bigram style comparison, to avoid pairing unrelated changes. When paired, one choice controls old-versus-current for that modification.",
                "a_cn": "展示层会判断相邻或相关的 removed 与 added 区块是否足够相似，能否视为一次替换。它使用规范化内容和相似度信号，包括字符二元组式比较，避免把无关变化强行配对。配对成功后，一个选择即可控制该修改使用旧版还是当前版。",
                "evidence": "Pairing logic and decision construction live in diffDisplay.js; the underlying source entries remain removed and added.",
                "pitfall": "Do not imply every adjacent delete/add pair becomes modified. / 不要暗示所有相邻删除和新增都会被合并成 modified。",
            },
            {
                "q_en": "How does the implementation control algorithmic cost?",
                "q_cn": "实现如何控制算法开销？",
                "a_en": "It does not run one unbounded quadratic comparison over the entire page. It aligns semantic blocks first, applies finer diff only where useful, and has explicit work caps: roughly 120k cells for page alignment, 80k for normal inline work, 240k for coarse sentence work, 40k for side-split inline work, and 500k for code-line work. Over-limit paths fall back to coarser but safe output.",
                "a_cn": "实现不会对整页执行一次无上限的二次复杂度比较。它先对齐语义区块，只在需要时做更细 diff，并设置明确工作量上限：页面对齐约 12 万格、普通行内约 8 万、粗粒度句子约 24 万、并排拆分行内约 4 万、代码行约 50 万。超限时退回更粗但安全的结果。",
                "evidence": "These caps are constants/guards in the utility and display code, protecting browser responsiveness and memory.",
                "pitfall": "Do not claim the full algorithm is linear. / 不要声称完整算法是线性复杂度。",
            },
            {
                "q_en": "Why not use a more sophisticated algorithm such as Myers diff everywhere?",
                "q_cn": "为什么不在所有地方都使用 Myers diff 等更复杂算法？",
                "a_en": "The main problem is not only shortest edit scripts; it is preserving Confluence structure and producing recoverable choices. Block extraction, canonicalisation, specialised structure handling, and safe Storage reconstruction matter more than swapping one sequence algorithm. LCS is simple to reason about, test, and cap for the Milestone 2 scope.",
                "a_cn": "核心问题不只是寻找最短编辑脚本，而是保留 Confluence 结构并生成可恢复的选择。区块提取、规范化、专门结构处理和安全 Storage 重建，比替换某一种序列算法更关键。对 Milestone 2 而言，LCS 易于解释、测试和设置上限。",
                "evidence": "The architecture separates block alignment from inline/table/layout-specific logic, so the sequence algorithm is only one layer.",
                "pitfall": "Do not argue that LCS is universally better; present it as a scoped engineering trade-off. / 不要说 LCS 普遍更优，应强调它是范围内的工程权衡。",
            },
            {
                "q_en": "How are stable selection keys maintained across views?",
                "q_cn": "如何在不同视图之间保持稳定的选择键？",
                "a_en": "Decision keys are derived from the canonical diff model rather than from screen position in one renderer. Inline and side-by-side views consume the same decision state, so changing the presentation mode does not reset which version the user selected for a changed unit.",
                "a_cn": "决策键来自规范差异模型，而不是某个渲染器中的屏幕位置。行内视图和并排视图共用同一份决策状态，因此切换展示模式不会重置用户对变化单元选择的版本。",
                "evidence": "diffDisplay.js builds shared choices consumed by both render modes; tests verify stable behaviour across view changes.",
                "pitfall": "Do not say the two views calculate independent diffs. / 不要说两个视图各自独立计算 diff。",
            },
        ],
    ),
    (
        "3. Rich text and special structures / 富文本与特殊结构",
        [
            {
                "q_en": "Why are rendered HTML and Storage XHTML kept as two tracks?",
                "q_cn": "为什么要同时保留 rendered HTML 和 Storage XHTML 两条轨道？",
                "a_en": "Rendered HTML is convenient for safe browser display and visual comparison, while Confluence Storage XHTML contains the authoritative macro, attribute, and structural data required for accurate write-back. The app sanitises and uses rendered content for presentation but retains raw Storage fragments for reconstruction.",
                "a_cn": "Rendered HTML 适合在浏览器中安全展示和视觉比较，而 Confluence Storage XHTML 包含准确写回所需的宏、属性和结构数据。应用会清理并使用 rendered 内容进行展示，同时保留原始 Storage 片段用于重建。",
                "evidence": "prepareConfluenceHtml handles display preparation; raw storage fragments travel separately into recovery composition.",
                "pitfall": "Do not publish sanitised preview HTML as if it were authoritative Storage. / 不要把清理后的预览 HTML 当作权威 Storage 发布。",
            },
            {
                "q_en": "What does the sanitisation step do?",
                "q_cn": "清理（sanitisation）步骤做什么？",
                "a_en": "It transforms Confluence-rendered HTML into a controlled preview representation, removing or neutralising unsafe or irrelevant elements and standardising content needed by the diff UI. This protects the host UI and improves deterministic comparison. Sanitisation is for display safety, not a substitute for Storage reconstruction.",
                "a_cn": "它把 Confluence 渲染 HTML 转换为受控的预览表示，移除或中和不安全、无关元素，并标准化 diff UI 所需内容。这既保护宿主界面，也提升比较的确定性。清理服务于展示安全，不能替代 Storage 重建。",
                "evidence": "prepareConfluenceHtml is implemented around utils.js:2850–3207 and has dedicated sanitisation tests.",
                "pitfall": "Do not equate sanitisation with full semantic parsing. / 不要把清理等同于完整语义解析。",
            },
            {
                "q_en": "How are blank lines represented and compared?",
                "q_cn": "空行如何表示和比较？",
                "a_en": "Blank lines are treated as countable structural content instead of being collapsed away. Consecutive blank runs are extracted and compared by count, so changing from two blank lines to five creates three added blanks. This preserves document spacing as an intentional edit.",
                "a_cn": "空行被视为可计数的结构内容，而不是直接折叠丢弃。连续空行会被提取并按数量比较，因此从 2 个空行变为 5 个空行会产生 3 个新增空行。这让文档间距也能被视作有意编辑。",
                "evidence": "Blank-line extraction and count-aware logic are within utils.js:3803–4192 and covered by blank-line tests.",
                "pitfall": "Do not say whitespace is always ignored. / 不要说所有空白都会被忽略。",
            },
            {
                "q_en": "How are soft breaks and paragraph boundaries handled?",
                "q_cn": "软换行和段落边界如何处理？",
                "a_en": "The implementation distinguishes meaningful block boundaries from inline breaks and normalises browser-specific HTML where needed. A Shift+Enter-style soft break should remain inside the relevant block, while separate paragraphs and explicit blank blocks remain independently alignable. This reduces false paragraph splits.",
                "a_cn": "实现会区分有意义的区块边界与行内换行，并在需要时规范化浏览器特定 HTML。类似 Shift+Enter 的软换行应保留在所属区块内部，而独立段落和显式空白区块仍可分别对齐，从而减少错误拆段。",
                "evidence": "HTML preparation plus semantic block extraction cooperate; the repository includes edge-case tests around line and blank handling.",
                "pitfall": "Do not promise identical DOM output from every Confluence renderer version. / 不要承诺每个 Confluence 渲染器版本都会产生完全相同的 DOM。",
            },
            {
                "q_en": "How are lists handled without losing structure?",
                "q_cn": "列表如何在不丢失结构的情况下处理？",
                "a_en": "Lists are decomposed into meaningful items for comparison while retaining enough wrapper information to reconstruct valid ordered or unordered list markup. Item-level choices improve usability, but recovery rebuilds the surrounding list containers rather than publishing isolated list-item fragments.",
                "a_cn": "列表会被拆解为有意义的条目进行比较，同时保留足够的外层信息，以重建有效的有序或无序列表标记。条目级选择提升可用性，但恢复时会重建列表容器，而不是发布孤立的列表项片段。",
                "evidence": "Semantic block logic contains list-aware branches and recovery composition groups selected list items under reconstructed wrappers.",
                "pitfall": "Do not claim every nested-list transformation is recovered at arbitrary depth with perfect minimal edits. / 不要声称任意深度的嵌套列表变换都能完美最小化恢复。",
            },
            {
                "q_en": "How are task lists treated?",
                "q_cn": "任务列表如何处理？",
                "a_en": "Task items are exposed as item-level changes so users can recover a specific task or status-bearing item. During write-back, the implementation reconstructs the task-list wrapper and preserves relevant Storage attributes instead of treating the rendered checkbox as the source of truth.",
                "a_cn": "任务条目以条目级变化呈现，因此用户可以恢复某个具体任务或带状态的条目。写回时，实现会重建任务列表外层并保留相关 Storage 属性，而不是把页面中渲染出来的复选框当作真实数据源。",
                "evidence": "Dedicated task handling exists in utility/recovery code and is validated in tasks-and-decisions tests.",
                "pitfall": "Do not say a task is just a normal bullet with a checkbox icon. / 不要说任务只是带复选框图标的普通项目符号。",
            },
            {
                "q_en": "How are decision lists treated?",
                "q_cn": "决策列表如何处理？",
                "a_en": "Decision items are compared and selected at item level, similar to tasks, because one decision may need recovery without replacing the whole group. The recovery path reconstructs the appropriate decision wrapper and uses Storage fragments to preserve Confluence semantics.",
                "a_cn": "决策条目与任务类似，按条目级进行比较和选择，因为用户可能只需要恢复某一项决策，而不是替换整组内容。恢复路径会重建相应决策外层，并使用 Storage 片段保留 Confluence 语义。",
                "evidence": "The same specialised test suite checks both task and decision display/recovery behaviour.",
                "pitfall": "Do not merge tasks and decisions conceptually; they share techniques but have different Storage semantics. / 不要在概念上把任务和决策合并，它们技术相似但 Storage 语义不同。",
            },
            {
                "q_en": "How are tables compared?",
                "q_cn": "表格如何比较？",
                "a_en": "The table path parses a logical grid rather than flattening the table into plain text. It can align compatible rows and cells, display cell-level changes, and preserve structural information such as spans. When structure becomes too incompatible or ambiguous, it falls back to a coarser whole-table change.",
                "a_cn": "表格路径会解析逻辑网格，而不是把整张表展平成纯文本。对于兼容结构，它能对齐行和单元格、展示单元格级变化，并保留合并等结构信息；当结构过于不兼容或存在歧义时，则退回更粗粒度的整表变化。",
                "evidence": "Table parsing, compatibility and rendering logic is concentrated in utils.js:5047–5893, with dedicated table tests.",
                "pitfall": "Do not promise spreadsheet-grade formula or semantic column analysis. / 不要承诺电子表格式的公式或语义列分析。",
            },
            {
                "q_en": "Can a user recover only one table cell?",
                "q_cn": "用户能否只恢复一个表格单元格？",
                "a_en": "The UI may highlight differences at cell level, but the current recovery guarantee is intentionally coarser: the selected recovery unit can be the whole table. This avoids reconstructing invalid combinations of row spans, column spans, headers, or table metadata.",
                "a_cn": "UI 可以在单元格层面高亮差异，但当前恢复保证有意保持更粗粒度：可恢复单元通常是整张表。这样可以避免把行合并、列合并、表头或表元数据组合成无效结构。",
                "evidence": "Display granularity and recovery granularity are deliberately different in the table implementation.",
                "pitfall": "Do not confuse cell-level visual diff with cell-level write-back. / 不要把单元格级视觉 diff 与单元格级写回混淆。",
            },
            {
                "q_en": "How are Confluence layouts handled?",
                "q_cn": "Confluence 布局如何处理？",
                "a_en": "Compatible layouts are decomposed into a boundary plus child content, allowing local differences to be compared without replacing the entire layout. Width or sizing changes can become a separate choice. If the old and current layouts are structurally incompatible, the system treats the layout as one larger replacement to protect validity.",
                "a_cn": "兼容布局会被拆成布局边界和子内容，从而比较局部差异，而不必替换整个布局。宽度或尺寸变化可以成为独立选择；若新旧布局结构不兼容，系统会把布局视为一个更大的替换单元，以保护结构有效性。",
                "evidence": "Layout-specific tests cover compatible children, incompatible layouts, width choices, and reconstruction.",
                "pitfall": "Do not say every layout mutation is independently recoverable. / 不要说所有布局变化都可以独立恢复。",
            },
            {
                "q_en": "How is code content diffed?",
                "q_cn": "代码内容如何进行 diff？",
                "a_en": "Code is handled with line-oriented logic so indentation and line boundaries remain meaningful. It has a larger but still bounded work cap—about 500k comparison cells—because code often contains many short lines. The renderer preserves code presentation while the recovery path keeps the authoritative source fragment.",
                "a_cn": "代码使用面向行的逻辑处理，使缩进和行边界保持意义。由于代码通常包含大量短行，它使用更大的但仍有限制的比较上限，约 50 万个比较单元。渲染器保留代码展示，恢复路径则保留权威源片段。",
                "evidence": "Code-line diff has its own bounded path rather than using the page-block cap unchanged.",
                "pitfall": "Do not claim language-aware AST diffing. / 不要声称实现了编程语言 AST 级 diff。",
            },
            {
                "q_en": "How are images, attachments, and mentions represented?",
                "q_cn": "图片、附件和用户提及如何表示？",
                "a_en": "The comparison uses stable identifiers and relevant attributes when available rather than relying only on rendered labels or URLs. The preview can show an understandable representation, while raw Storage preserves attachment, image, or mention metadata required by Confluence.",
                "a_cn": "比较会在可用时使用稳定标识符和相关属性，而不是只依赖渲染标签或 URL。预览可以展示易懂表示，原始 Storage 则保留 Confluence 所需的附件、图片或用户提及元数据。",
                "evidence": "Canonicalisation includes rich-content-specific normalisation; raw preservation protects attributes not visible in rendered HTML.",
                "pitfall": "Do not use a changing display name as the sole identity of a mention. / 不要把可能变化的显示名作为用户提及的唯一身份。",
            },
            {
                "q_en": "What happens to unsupported macros or raw blocks?",
                "q_cn": "不支持的宏或原始区块会怎样处理？",
                "a_en": "Unsupported content is not silently discarded. The system keeps raw Storage fragments and can expose the unit as a coarse change. If a trustworthy Storage representation is unavailable, write-back is disabled rather than guessing how to rebuild the content.",
                "a_cn": "不支持的内容不会被静默丢弃。系统会保留原始 Storage 片段，并可把该单元作为较粗粒度变化展示。如果无法取得可信的 Storage 表示，则禁用写回，而不是猜测如何重建内容。",
                "evidence": "Raw-block preservation and recovery guards are part of the fallback path; tests cover raw-storage availability conditions.",
                "pitfall": "Do not say unsupported means deleted or ignored. / 不要说“不支持”等同于删除或忽略。",
            },
            {
                "q_en": "How are formatting-only changes treated?",
                "q_cn": "仅格式变化如何处理？",
                "a_en": "Canonicalisation removes some irrelevant presentation noise, but meaningful inline formatting can remain visible through finer diff or raw-content comparison. The exact granularity depends on the structure and whether the change can be represented safely. The priority is a useful, recoverable result rather than highlighting every DOM-level mutation.",
                "a_cn": "规范化会去除部分无关展示噪声，但有意义的行内格式变化仍可通过细粒度 diff 或原始内容比较体现。具体粒度取决于结构以及变化是否能安全表示。系统优先提供有用、可恢复的结果，而不是高亮每一个 DOM 层变化。",
                "evidence": "The pipeline separates canonical matching, display HTML, and raw Storage, allowing different notions of equality at different stages.",
                "pitfall": "Do not promise pixel-perfect visual diffs. / 不要承诺像素级视觉 diff。",
            },
        ],
    ),
    (
        "4. User decisions, preview, and views / 用户决策、预览与视图",
        [
            {
                "q_en": "What does choosing ‘current’ or ‘old’ actually mean?",
                "q_cn": "选择 current 或 old 实际意味着什么？",
                "a_en": "For each recoverable changed unit, current means keep the live page’s contribution and old means take the historical contribution. Unchanged units pass through automatically. The final preview is the ordered composition of all unchanged content plus each decision’s selected contribution.",
                "a_cn": "对每个可恢复变化单元，current 表示保留线上页面贡献，old 表示采用历史版本贡献；未变化单元自动通过。最终预览由所有未变化内容与每项决策所选贡献按顺序组合而成。",
                "evidence": "Choice state feeds preview composition and later Storage reconstruction; it is not merely a visual highlight state.",
                "pitfall": "Do not describe old/current as left/right screen columns only. / 不要把 old/current 仅描述成屏幕左右列。",
            },
            {
                "q_en": "Why is current the default choice?",
                "q_cn": "为什么默认选择 current？",
                "a_en": "Current is the safer default because opening the tool should not implicitly undo live work. Historical content is restored only through an explicit user choice. This supports least-surprise behaviour and reduces the risk of accidental rollback.",
                "a_cn": "current 是更安全的默认值，因为打开工具不应自动撤销线上工作。只有用户明确选择时才恢复历史内容，这符合最小意外原则，也降低误回滚风险。",
                "evidence": "Decision initialisation defaults changed units to current; bulk restore is an explicit action.",
                "pitfall": "Do not say the system automatically selects the older side when it sees a deletion. / 不要说系统发现删除后会自动选择旧版。",
            },
            {
                "q_en": "What is the purpose of inline and side-by-side views?",
                "q_cn": "行内视图和并排视图分别有什么用途？",
                "a_en": "Inline view is compact and good for reading the document as one narrative with additions and removals marked in place. Side-by-side view is better for comparing old and current context directly. Both are projections of the same canonical diff and the same recovery decisions.",
                "a_cn": "行内视图更紧凑，适合把文档作为一条连续叙事阅读，并在原位标记新增与删除；并排视图更适合直接比较新旧上下文。两者都是同一份规范 diff 和同一组恢复决策的不同投影。",
                "evidence": "Shared choice keys prevent view switching from changing the composed preview.",
                "pitfall": "Do not present the two views as different recovery modes. / 不要把两个视图说成不同恢复模式。",
            },
            {
                "q_en": "How does the side-by-side view align content?",
                "q_cn": "并排视图如何对齐内容？",
                "a_en": "It starts from the ordered canonical diff stream, then pairs related removed and added blocks and inserts empty counterparts where one side has no matching block. Finer inline markers can be split between the two sides under a bounded workload, improving correspondence without recalculating the whole document.",
                "a_cn": "它从有序的规范差异流出发，配对相关的 removed 与 added 区块，并在某一侧没有对应内容时插入空位。在有限工作量内，还可以把更细的行内标记拆分到左右两侧，以提升对应关系，而无需重新计算整份文档。",
                "evidence": "The side-split inline path has an approximately 40k-cell cap and uses shared display decisions.",
                "pitfall": "Do not promise perfect vertical alignment for arbitrarily large or nested blocks. / 不要承诺任意大型或嵌套区块都能像素级垂直对齐。",
            },
            {
                "q_en": "What do bulk actions do?",
                "q_cn": "批量操作做什么？",
                "a_en": "Bulk actions update the decision state for many eligible changed units, for example keeping all current content or restoring all historical content. They do not bypass preview or safety checks. Users can still refine individual decisions before publishing.",
                "a_cn": "批量操作会更新多个符合条件的变化单元决策，例如全部保留当前内容或全部恢复历史内容。它们不会绕过预览和安全检查，用户仍可在发布前逐项调整。",
                "evidence": "Bulk commands manipulate the same per-choice state used by individual controls and preview composition.",
                "pitfall": "Do not call ‘restore all’ an immediate publish action. / 不要把“全部恢复”说成会立即发布。",
            },
            {
                "q_en": "What does the preview guarantee?",
                "q_cn": "预览提供什么保证？",
                "a_en": "The preview shows the expected composed content under the current decisions and helps catch obvious mistakes before write-back. It is not an immutable guarantee that the live page has not changed since loading. That separate concurrency condition is checked again immediately before publishing.",
                "a_cn": "预览展示当前决策下预期组合出的内容，帮助用户在写回前发现明显错误。但它不能保证页面自加载以来一直没有变化；这个并发条件会在发布前再次独立检查。",
                "evidence": "Preview composition and publish-time version validation are separate stages.",
                "pitfall": "Do not claim visual preview alone prevents stale writes. / 不要声称仅靠视觉预览就能防止过期写入。",
            },
            {
                "q_en": "What are Difference Notes for?",
                "q_cn": "Difference Notes 有什么作用？",
                "a_en": "Difference Notes provide a concise human-readable summary of the selected recovery result or relevant changes. They support auditability and user confidence, but they are not the authoritative reconstruction format. Storage XHTML remains the write-back source.",
                "a_cn": "Difference Notes 用简洁、可读的方式总结所选恢复结果或相关变化，帮助审计并增强用户信心，但它不是权威重建格式；写回仍以 Storage XHTML 为准。",
                "evidence": "Notes are generated alongside the preview rather than replacing the structured composition path.",
                "pitfall": "Do not describe notes as a patch file applied to Confluence. / 不要把 Notes 说成直接应用到 Confluence 的补丁文件。",
            },
            {
                "q_en": "What happens if the user selects another historical version?",
                "q_cn": "如果用户切换到另一个历史版本，会发生什么？",
                "a_en": "The comparison basis changes, so the diff and its decision set must be rebuilt for the newly selected version. Previous choices should not be blindly reused because their keys and semantics may no longer refer to the same content. The UI resets or regenerates state accordingly.",
                "a_cn": "比较基准改变后，必须针对新选择的历史版本重新构建 diff 和决策集合。不能盲目沿用旧选择，因为原有键和语义可能不再指向同一内容；UI 会相应重置或重新生成状态。",
                "evidence": "Version-selection state is upstream of diff construction and choice initialisation in TextDiffPage.",
                "pitfall": "Do not promise selections persist across unrelated version comparisons. / 不要承诺选择会跨不相关版本比较持续存在。",
            },
        ],
    ),
    (
        "5. Recovery safety and Forge architecture / 恢复安全与 Forge 架构",
        [
            {
                "q_en": "How is the final Confluence Storage content reconstructed?",
                "q_cn": "最终的 Confluence Storage 内容如何重建？",
                "a_en": "The app walks the ordered diff decisions and chooses authoritative old or current Storage fragments. It groups structures that require wrappers—such as lists, tasks, decisions, tables, and layouts—and reconstructs valid containers. Unchanged and unsupported raw fragments are preserved rather than regenerated from visible text.",
                "a_cn": "应用按顺序遍历差异决策，为每个单元选择权威的旧版或当前版 Storage 片段。对于列表、任务、决策、表格和布局等需要外层结构的内容，它会分组并重建有效容器。未变化和不支持的原始片段会被保留，而不是从可见文本重新生成。",
                "evidence": "Recovery is driven by raw fragments attached to diff entries; specialised grouping handles wrapper-based structures.",
                "pitfall": "Do not say the system converts the rendered DOM back into Storage generically. / 不要说系统会通用地把渲染 DOM 转回 Storage。",
            },
            {
                "q_en": "How does the app prevent overwriting newer edits?",
                "q_cn": "应用如何防止覆盖更新的编辑？",
                "a_en": "Immediately before update, the backend re-reads the live page using the user context and compares its version number with the expected version used for the diff. If it has changed, the save is rejected and the user must refresh or recompute. A successful update writes the next version number.",
                "a_cn": "更新前，后端会以用户身份重新读取线上页面，并把其版本号与生成 diff 时的预期版本比较。如果版本已变化，则拒绝保存，要求用户刷新或重新计算；只有校验成功时才写入下一个版本号。",
                "evidence": "The resolver’s write path performs an expectedVersionNumber check and sends version + 1 on update.",
                "pitfall": "Do not call this a merge of concurrent edits; it is conflict detection and rejection. / 不要把它说成并发编辑自动合并；它是冲突检测与拒绝。",
            },
            {
                "q_en": "Why is asUser used for Confluence REST requests?",
                "q_cn": "为什么 Confluence REST 请求使用 asUser？",
                "a_en": "asUser applies the current user’s Confluence permissions to the operation. That means the app does not gain broader editing power than the user and product authorisation remains part of the request. It is the safer default for user-initiated reads and writes.",
                "a_cn": "asUser 会把当前用户的 Confluence 权限应用到操作中，因此应用不会获得超出用户自身的编辑能力，产品授权检查也仍属于请求的一部分。对于用户发起的读写，这是更安全的默认方式。",
                "evidence": "Resolver REST calls use api.asUser() rather than relying on app-wide authority for user actions.",
                "pitfall": "Do not say asUser removes the need for scopes or all validation. / 不要说 asUser 可以取消 scope 或所有业务校验。",
            },
            {
                "q_en": "What happens if raw Storage content is missing?",
                "q_cn": "如果缺少原始 Storage 内容会怎样？",
                "a_en": "The system can still show limited information from rendered content, but it should not offer a write-back that requires guessing the Storage representation. Recovery is disabled or constrained until authoritative raw content is available. This is a deliberate fail-safe.",
                "a_cn": "系统仍可利用 rendered 内容显示有限信息，但不应提供需要猜测 Storage 表示的写回操作。只有获得权威原始内容后才允许相应恢复，否则会禁用或限制恢复。这是有意的故障安全设计。",
                "evidence": "Write enablement is guarded by storage availability and valid reconstruction inputs.",
                "pitfall": "Do not fabricate Storage from the preview just to keep the button enabled. / 不要为了让按钮可用而从预览中臆造 Storage。",
            },
            {
                "q_en": "Is there a payload-size safeguard?",
                "q_cn": "是否有请求体大小保护？",
                "a_en": "Yes. The write path enforces an approximate two-megabyte payload limit before attempting the page update. This protects the function and downstream API from unexpectedly large reconstructed content and gives a controlled error instead of an unstable request.",
                "a_cn": "有。写回路径在尝试更新页面前执行约 2 MB 的请求体大小限制，以保护函数和下游 API，避免超大重建内容造成不稳定请求，并返回可控错误。",
                "evidence": "The resolver checks payload length before issuing the update request.",
                "pitfall": "Do not claim 2 MB is Confluence’s universal page limit; it is this app’s defensive guard. / 不要说 2 MB 是 Confluence 的统一页面上限；这是本应用的防御性保护。",
            },
            {
                "q_en": "Why use a Forge resolver for write-back instead of only frontend requests?",
                "q_cn": "为什么写回使用 Forge resolver，而不完全放在前端？",
                "a_en": "The resolver centralises version validation, payload validation, error normalisation, and the final privileged product request. It creates one auditable boundary for a high-impact action. Read-only UI requests can remain simpler, but publishing benefits from server-side control.",
                "a_cn": "resolver 集中处理版本校验、请求体校验、错误规范化和最终产品请求，为高影响操作建立一个可审计边界。只读 UI 请求可以更简单，但发布操作更适合由后端集中控制。",
                "evidence": "The frontend sends the composed Storage plus expected version; the resolver rechecks and performs the update.",
                "pitfall": "Do not say a resolver is automatically secure; its checks and asUser usage make the path safer. / 不要说用了 resolver 就自动安全；真正起作用的是校验与 asUser。",
            },
            {
                "q_en": "What permissions and scopes should the app use?",
                "q_cn": "应用应使用哪些权限和 scopes？",
                "a_en": "Only the minimum scopes needed to read page versions/content and update content should be declared. User-context product requests then enforce the user’s own access. Any future feature that adds a new API should justify and lint the additional scope before deployment.",
                "a_cn": "应用只应声明读取页面版本/内容以及更新内容所需的最小 scopes，并通过用户上下文产品请求继续执行用户自身权限。未来功能若增加新 API，应先证明新增 scope 的必要性并通过 lint，再进行部署。",
                "evidence": "The repository manifest is the authority for scopes; Forge lint should be used after manifest changes.",
                "pitfall": "Do not request broad scopes ‘for future use’. / 不要为了“以后可能用到”而申请宽泛 scope。",
            },
            {
                "q_en": "What error behaviour should the tutor expect to see?",
                "q_cn": "tutor 应该看到怎样的错误处理行为？",
                "a_en": "Expected failures—missing versions, unavailable Storage, stale version numbers, oversized payloads, and REST errors—should produce understandable UI feedback and avoid partial write-back. The app should preserve the user’s context where possible so they can refresh or adjust rather than losing their work.",
                "a_cn": "对于版本缺失、Storage 不可用、版本号过期、请求体过大和 REST 错误等预期失败，UI 应给出可理解反馈，并避免部分写回。在可能情况下保留用户上下文，让用户可以刷新或调整，而不是丢失操作。",
                "evidence": "TextDiffPage handles loading/error states; resolver error paths reject the update before a successful publish response.",
                "pitfall": "Do not demo a forced error unless you have prepared a recoverable path. / 不要在未准备恢复路径时现场强行制造错误。",
            },
        ],
    ),
    (
        "6. Testing, limits, and future work / 测试、限制与后续计划",
        [
            {
                "q_en": "What automated test evidence supports Milestone 2?",
                "q_cn": "有哪些自动化测试证据支撑 Milestone 2？",
                "a_en": "On the latest pulled code, the complete frontend test run passes 11 suites and 169 tests. Coverage is organised around utility behaviour and page-level interaction, including sanitisation, inline diff, tables, layouts, tasks/decisions, blank lines, side-by-side rendering, and selective recovery.",
                "a_cn": "在最新拉取的代码上，完整前端测试运行通过 11 个测试套件、169 个测试。覆盖按工具逻辑和页面级交互组织，包括清理、行内 diff、表格、布局、任务/决策、空行、并排渲染和选择性恢复。",
                "evidence": "Local verification was performed at commit 2e3e05718c37968df25642b68c2c82db39ebfde4.",
                "pitfall": "Do not turn passing tests into a claim of zero production bugs. / 不要把测试全通过说成生产环境零缺陷。",
            },
            {
                "q_en": "Which tests are most important to explain in the presentation?",
                "q_cn": "presentation 中最值得解释哪些测试？",
                "a_en": "Prioritise tests that prove the project’s risky claims: stable old/current decisions, blank-line counts, table compatibility fallback, layout reconstruction, task/decision wrappers, side-by-side consistency, sanitisation, and stale-write prevention. These show that the tests validate semantics, not only component rendering.",
                "a_cn": "应优先解释能证明高风险主张的测试：稳定的新旧选择、空行计数、表格兼容性回退、布局重建、任务/决策外层、并排视图一致性、清理和过期写入防护。这些测试说明我们验证的是语义，而不只是组件能否渲染。",
                "evidence": "The repository names dedicated suites around these structures; TextDiffPage tests exercise interaction and recovery state.",
                "pitfall": "Do not spend most of the answer listing trivial snapshot tests. / 不要把回答重点放在琐碎快照测试上。",
            },
            {
                "q_en": "What is the biggest current technical limitation?",
                "q_cn": "当前最大的技术限制是什么？",
                "a_en": "There is an unavoidable trade-off between recovery granularity and structural validity. Complex or incompatible structures—especially tables, layouts, nested content, and unknown macros—may fall back to a larger recovery unit. The system prioritises preserving valid Confluence Storage over offering an unsafe microscopic edit.",
                "a_cn": "当前最重要的限制是恢复粒度与结构有效性之间的权衡。复杂或不兼容结构，尤其是表格、布局、嵌套内容和未知宏，可能退回更大的恢复单元。系统优先保证 Confluence Storage 有效，而不是提供不安全的微观编辑。",
                "evidence": "Whole-table and whole-layout fallbacks, plus raw-block preservation, embody this design decision.",
                "pitfall": "Do not hide the limitation; explain why the fallback is safer. / 不要回避限制，应解释为什么回退更安全。",
            },
            {
                "q_en": "How does the tool behave on very large pages?",
                "q_cn": "工具在超大页面上如何表现？",
                "a_en": "The browser work is bounded by per-stage comparison caps, and expensive fine-grained paths degrade to coarser output. This keeps the interface more responsive and avoids uncontrolled matrix allocation. However, very large or structure-heavy pages remain a performance area for profiling and future optimisation.",
                "a_cn": "浏览器端工作量由各阶段比较上限控制，昂贵的细粒度路径会降级为更粗结果，从而保持界面响应并避免无控制的矩阵分配。不过，超大或结构密集页面仍是需要持续性能分析和优化的领域。",
                "evidence": "Page, inline, sentence, side-split, and code paths each have explicit workload guards.",
                "pitfall": "Do not claim the caps make performance constant-time. / 不要声称设置上限后性能就变成常数时间。",
            },
            {
                "q_en": "What would you improve in the next milestone?",
                "q_cn": "下一个 milestone 你会优先改进什么？",
                "a_en": "I would prioritise production telemetry and profiling, broader fixture coverage using real Confluence Storage examples, clearer explanations for coarse fallbacks, and carefully expanding recoverable granularity for tables/layouts where validity can be proven. I would also reduce build-tool technical debt and add stronger end-to-end validation in a real Forge environment.",
                "a_cn": "我会优先加入生产遥测与性能分析，使用真实 Confluence Storage 样例扩大 fixture 覆盖，为粗粒度回退提供更清晰解释，并在能证明结构有效的前提下逐步提高表格/布局恢复粒度。同时减少构建工具技术债，并在真实 Forge 环境中加强端到端验证。",
                "evidence": "The current strengths are unit/integration coverage and defensive reconstruction; the natural next step is production evidence and broader real-content fixtures.",
                "pitfall": "Do not promise a complete rewrite. Propose incremental improvements around measured risks. / 不要承诺完全重写，应围绕已测量风险渐进改进。",
            },
            {
                "q_en": "Why not use an AI model to decide which changes to restore?",
                "q_cn": "为什么不使用 AI 模型自动决定恢复哪些变化？",
                "a_en": "Recovery is a high-impact authoring action, and intent is contextual. The current design keeps the user in control and uses deterministic rules for diff and reconstruction. AI could later assist with summaries or recommendations, but it should not silently choose or publish historical content without explicit confirmation.",
                "a_cn": "恢复是高影响的内容编辑操作，而用户意图高度依赖上下文。当前设计让用户保持控制，并以确定性规则完成 diff 与重建。未来 AI 可以辅助总结或推荐，但不应在没有明确确认时静默选择或发布历史内容。",
                "evidence": "All changed units default to current and require explicit decisions; publishing remains a separate confirmed action.",
                "pitfall": "Do not present ‘no AI’ as a weakness; frame deterministic recovery as a trust and safety choice. / 不要把“没有 AI”说成缺点，应说明确定性恢复是信任与安全选择。",
            },
            {
                "q_en": "How would you demonstrate the feature reliably in three minutes?",
                "q_cn": "你会如何在三分钟内稳定演示该功能？",
                "a_en": "Use a prepared page with one paragraph edit, one deleted list/task item, one blank-line change, and one small table edit. Select a known historical version, explain old-to-current direction, show inline and side-by-side views, restore two specific units, verify the preview, then publish only if the environment is stable. Keep screenshots as a fallback.",
                "a_cn": "使用预先准备的页面，包含一处段落修改、一处列表/任务项删除、一处空行变化和一处小表格修改。选择已知历史版本，解释 old-to-current 方向，展示行内与并排视图，恢复两个指定单元，检查预览；只有环境稳定时才发布，并准备截图作为备用。",
                "evidence": "This path demonstrates the core model, specialised structures, shared decisions, preview, and safety without relying on rare edge cases.",
                "pitfall": "Do not start with the largest page or an unknown macro-heavy page. / 不要用最大页面或充满未知宏的页面开场。",
            },
            {
                "q_en": "If a tutor says ‘this is just a diff library’, how do you respond?",
                "q_cn": "如果 tutor 说“这不就是一个 diff 库吗”，你怎么回答？",
                "a_en": "The sequence diff is only one component. The engineering value lies in adapting Confluence rich-text representations into semantic blocks, preserving unsupported Storage, providing structure-aware display and recovery choices, reconstructing valid XHTML, and preventing stale or unauthorised writes. A generic diff library does not solve those product and integration problems.",
                "a_cn": "序列 diff 只是其中一个组件。工程价值在于把 Confluence 富文本表示转换成语义区块、保留未支持 Storage、提供结构感知展示与恢复选择、重建有效 XHTML，并防止过期或未授权写入。通用 diff 库不能直接解决这些产品与集成问题。",
                "evidence": "The codebase contains dedicated Confluence sanitisation, canonicalisation, table/layout/task handling, decision modelling, and resolver safety beyond LCS.",
                "pitfall": "Do not dismiss libraries; explain the application-specific layers built around the algorithm. / 不要贬低库，应解释围绕算法实现的应用专用层。",
            },
            {
                "q_en": "What is the strongest one-sentence conclusion for the tutor?",
                "q_cn": "面对 tutor，最有力的一句总结是什么？",
                "a_en": "Milestone 2 proves that selective Confluence history recovery can be made understandable and safe by separating display from Storage, diffing semantic blocks, keeping users in control, and validating every write against the live page.",
                "a_cn": "Milestone 2 证明了：通过分离展示与 Storage、比较语义区块、让用户保持控制并在写回前校验线上页面，可以让 Confluence 历史内容的选择性恢复既易懂又安全。",
                "evidence": "This sentence connects the product value, technical architecture, user-control model, and safety boundary.",
                "pitfall": "Do not end with only ‘we built a text diff’. State the recovery value and safety model. / 不要只用“我们做了 text diff”结尾，要说出恢复价值和安全模型。",
            },
        ],
    ),
]


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MILESTONE 2")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Text Diff Deep Q&A Template")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("文本差异深度双语问答模板")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    rr.font.color.rgb = RGBColor.from_string(YELLOW)
    rr.font.size = Pt(13)

    add_callout(
        doc,
        "How to use / 使用方式",
        "Memorise the Chinese answer first, then practise the English answer aloud. Use the code-evidence line only when the tutor asks a follow-up. / 先熟记中文回答，再练习英文口述；只有 tutor 追问时，再补充代码证据。",
        LIGHT_GREY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.add_run("Project: Dynamic History for Confluence\n").bold = True
    p.add_run("Code baseline: 2e3e05718c37968df25642b68c2c82db39ebfde4\n")
    p.add_run("Prepared: 22 July 2026 · Australia/Sydney")
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(MID_GREY)

    doc.add_page_break()


def add_quick_reference(doc: Document) -> None:
    doc.add_heading("Fast preparation map / 快速准备地图", level=1)
    add_callout(
        doc,
        "Core framing / 核心表述",
        "This is not a plain-text highlighter. It is a structure-aware comparison and selective-recovery system that separates safe display from authoritative Confluence Storage. / 这不是纯文本高亮器，而是一个分离安全展示与权威 Confluence Storage 的结构感知比较和选择性恢复系统。",
    )
    doc.add_heading("Ten answers to memorise first / 优先背熟的十个答案", level=2)
    items = [
        "Old → current direction: removed existed before; added exists now. / 方向是旧版到当前版：removed 过去存在，added 当前存在。",
        "Main primitives are same, removed, added; modified is a display pairing. / 主原语是 same、removed、added；modified 是展示层配对。",
        "Semantic blocks protect structure better than one plain-text string. / 语义区块比整段纯文本更能保护结构。",
        "LCS aligns ordered blocks; explicit caps prevent uncontrolled quadratic work. / LCS 对齐有序区块；显式上限控制二次复杂度工作量。",
        "Rendered HTML is for display; Storage XHTML is for recovery and publish. / Rendered HTML 用于展示；Storage XHTML 用于恢复和发布。",
        "Current is the safe default; restoring old content requires explicit choice. / current 是安全默认值；恢复旧内容必须明确选择。",
        "Cell-level table highlighting does not imply cell-level write-back. / 表格单元格级高亮不等于单元格级写回。",
        "The backend re-reads the page and rejects stale-version writes. / 后端重新读取页面并拒绝过期版本写入。",
        "Unsupported raw content is preserved; missing authoritative Storage disables write-back. / 未支持原始内容会保留；缺少权威 Storage 时禁用写回。",
        "Latest local evidence: 11 suites and 169 tests pass. / 最新本地证据：11 个测试套件、169 个测试全部通过。",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Answer pattern / 回答结构", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ("Step / 步骤", "What to say / 应说内容")):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.add_run(text).bold = True
    rows = [
        ("1. Outcome / 结论", "State the product value in one sentence. / 用一句话说明产品价值。"),
        ("2. Mechanism / 机制", "Name the relevant pipeline or safety mechanism. / 点出相关管线或安全机制。"),
        ("3. Evidence / 证据", "Use a code area, test result, or explicit limit. / 使用代码区域、测试结果或明确上限。"),
        ("4. Boundary / 边界", "Admit the current granularity or fallback. / 主动说明当前粒度或回退边界。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        set_table_geometry(table, [2500, 6860])
    doc.add_page_break()


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    for section in doc.sections:
        set_repeat_header_text(section, "DYNAMIC HISTORY · MILESTONE 2 · TEXT DIFF Q&A")
        add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    add_quick_reference(doc)

    number = 1
    for section_title, questions in SECTIONS:
        doc.add_heading(section_title, level=1)
        if number == 1:
            add_callout(
                doc,
                "Presentation rule / 答辩规则",
                "Lead with the product outcome, then explain one mechanism and one limitation. / 先讲产品结果，再讲一个机制和一个限制。",
            )
        for question in questions:
            add_question(doc, number, question)
            number += 1

    doc.add_heading("Final rehearsal checklist / 最终练习清单", level=1)
    checklist = [
        "I can explain historical → current direction without looking at notes. / 我能不看笔记解释历史版到当前版的方向。",
        "I can distinguish formal diff types from display decisions. / 我能区分正式 diff 类型和展示决策。",
        "I can explain why preview HTML and Storage XHTML are separate. / 我能解释预览 HTML 和 Storage XHTML 为何分离。",
        "I can state one performance cap and one safe fallback. / 我能说出一个性能上限和一个安全回退。",
        "I can explain the stale-version protection before publish. / 我能解释发布前的过期版本防护。",
        "I will describe limitations honestly and connect them to structural validity. / 我会坦诚说明限制，并把它与结构有效性联系起来。",
        "I have a prepared demo page and screenshots as fallback. / 我准备了演示页面和备用截图。",
    ]
    for item in checklist:
        doc.add_paragraph("☐ " + item)

    add_callout(
        doc,
        "Closing line / 收尾句",
        "Milestone 2 proves that selective Confluence history recovery can be understandable, structure-aware, and safe. / Milestone 2 证明了 Confluence 历史内容的选择性恢复可以做到易懂、结构感知且安全。",
        LIGHT_GREY,
    )

    props = doc.core_properties
    props.title = "Milestone 2 Text Diff Deep Q&A Template — Bilingual"
    props.subject = "Dynamic History for Confluence presentation preparation"
    props.author = "Codex"
    props.keywords = "Milestone 2, text diff, Confluence, bilingual, tutor Q&A"
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")
    print(f"Questions: {number - 1}")
    return OUTPUT


if __name__ == "__main__":
    build()
