from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("outputs/tutor-speech-script-bilingual.docx")


def set_run_font(run, size, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.2):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, 22, bold=True, color=RGBColor(0x0B, 0x25, 0x45))


def add_subtitle(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=14, line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, 11, color=RGBColor(0x55, 0x55, 0x55))


def add_heading(document, text, level=1):
    sizes = {1: 16, 2: 13}
    colors = {1: RGBColor(0x2E, 0x74, 0xB5), 2: RGBColor(0x1F, 0x4D, 0x78)}
    before = {1: 16, 2: 10}
    after = {1: 8, 2: 5}

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=before[level], after=after[level], line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, sizes[level], bold=True, color=colors[level])


def add_body(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, before=0, after=4, line=1.2)
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=10.5, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Tutor speech script | English + Chinese")
    set_run_font(run, 9, color=RGBColor(0x66, 0x66, 0x66))


def add_summary_table(document):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(0.9)
    table.columns[1].width = Inches(2.6)
    table.columns[2].width = Inches(3.0)

    headers = ["页码", "主题", "建议时长"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color=RGBColor(0x1F, 0x4D, 0x78))

    rows = [
        ["1", "项目定位", "25 到 30 秒"],
        ["2", "五层处理流程", "40 到 45 秒"],
        ["3", "不同内容类型的 diff 策略", "45 到 50 秒"],
        ["4", "工程配套操作", "30 到 35 秒"],
        ["5", "总结与价值", "20 到 30 秒"],
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def add_slide_section(document, slide_number, slide_title, english_script, chinese_translation):
    add_heading(document, f"Slide {slide_number}: {slide_title}", level=1)

    add_heading(document, "English Script", level=2)
    add_body(document, english_script)

    add_heading(document, "中文翻译", level=2)
    add_body(document, chinese_translation)


def build_document():
    document = Document()
    configure_page(document)
    add_footer(document.sections[0])

    add_title(document, "Tutor 演示英文讲稿与中文翻译")
    add_subtitle(
        document,
        "对应演示文件：tutor-text-diff-presentation-en.pptx | 用途：方便口头练习与临场准备"
    )

    add_body(
        document,
        "这份文档按照英文 PPT 的页顺序整理了完整讲稿。每一页都先给出可直接朗读的英文版本，再附上中文翻译，方便你对照理解和背诵。整体语气偏自然、简洁，适合 tutor 演示场景。"
    )

    add_heading(document, "使用建议", level=1)
    add_bullets(
        document,
        [
            "先按英文完整读一遍，再看中文确认含义。",
            "如果时间紧，可以优先记每页开头和结尾句。",
            "正式讲的时候，不需要一字不差背诵，保持逻辑顺序一致就可以。",
        ]
    )

    add_heading(document, "时长分配", level=1)
    add_summary_table(document)

    add_slide_section(
        document,
        1,
        "Text Diff in Dynamic History",
        "This project is about text diff in Dynamic History. The key idea is that we do not rely on one single diff function. Instead, we combine several steps into one practical pipeline. The input is Confluence storage HTML from an old version and the current version. Our goal is to produce a comparison that is readable, stable, and useful for the frontend. So from the beginning, this is not only an algorithm problem. It is also a product and engineering problem.",
        "这一页我主要介绍项目定位。这个项目讨论的是 Dynamic History 里的文本差异比较。核心思想是，我们并不是依赖某一个单独的 diff 函数，而是把多个步骤组合成一条实用的处理流水线。输入是旧版本和当前版本的 Confluence storage HTML，目标是生成一个既易读、又稳定、还能方便前端使用的比较结果。所以从一开始，这就不只是一个算法问题，也是一个产品和工程问题。"
    )

    add_slide_section(
        document,
        2,
        "Five processing layers",
        "On this slide, I explain the main pipeline in five layers. First, we preprocess the rich content, such as links, code macros, emoticons, and image attachments. Second, we split the page into comparable blocks like paragraphs, headings, lists, tables, and code blocks. Third, we use block-level LCS to align the old and current content. Fourth, we choose the right comparison mode for each block. Finally, we return rendered HTML, structured blocks, summary counters, and limited flags for the UI.",
        "这一页我解释主流程，一共分成五层。第一步是预处理富文本内容，比如链接、代码宏、表情和图片附件。第二步是把页面拆成可比较的块，例如段落、标题、列表、表格和代码块。第三步是使用块级 LCS 对旧内容和当前内容做对齐。第四步是为每一种块选择合适的比较方式。最后，我们返回可渲染的 HTML、结构化 blocks、summary 计数，以及 limited 标记，供前端界面使用。"
    )

    add_slide_section(
        document,
        3,
        "One project, several comparison modes",
        "The most important point here is that one diff strategy is not enough for all content types. Normal text uses inline token diff, because users want to see small additions and deletions inside a sentence. Very large text uses a fallback strategy, so the browser does not freeze. Code blocks are compared line by line, because indentation and line structure matter. Tables are handled differently again. If the table shape stays the same, we highlight changed cells. If the shape changes, we show the old and new tables side by side.",
        "这一页最重要的一点是，一种 diff 策略并不足以处理所有内容类型。普通文本使用行内 token diff，因为用户希望看到句子内部的小范围新增和删除。特别大的文本会使用回退策略，这样浏览器就不会卡住。代码块按行比较，因为缩进和行结构很重要。表格则采用另一套逻辑。如果表格结构没变，我们高亮变化的单元格；如果结构发生变化，我们就把旧表格和新表格并排展示。"
    )

    add_slide_section(
        document,
        4,
        "Supporting operations",
        "This slide shows why the surrounding engineering work is also important. The backend fetches page versions, attachment URLs, and author names. Then the frontend normalizes storage-format HTML before comparison. After that, the rendering logic keeps only safe tags and attributes, so risky content does not reach the preview. Finally, the UI has a fallback path. If diff rendering fails, the comparison panel shows a safe message instead of breaking the whole page. These supporting operations make the feature reliable in practice.",
        "这一页说明为什么周边工程工作同样重要。后端会先拉取页面版本、附件 URL 和作者信息。然后前端会在比较之前规范化 storage 格式的 HTML。接着，渲染逻辑只保留安全的标签和属性，避免有风险的内容进入预览。最后，界面本身也有兜底路径。如果 diff 渲染失败，比较面板会显示一个安全提示，而不是让整个页面崩溃。正是这些配套操作，让这个功能在实际使用中更可靠。"
    )

    add_slide_section(
        document,
        5,
        "Why this design is practical",
        "To conclude, I think this design is practical for four reasons. First, it is readable for users, because they can understand rich-text changes directly. Second, it is stable for large content, because expensive cases fall back to safer paths. Third, it is specialized for code blocks and tables, so important structure is preserved. And fourth, it is connected to real product behavior, including fetching, sanitization, rendering, summary, and fallback. In one sentence, this project turns a classic LCS-style idea into a usable rich-text diff feature.",
        "最后总结一下，我认为这套设计是实用的，主要有四个原因。第一，它对用户来说更易读，因为用户可以直接理解富文本里的变化。第二，它面对大内容时更稳定，因为代价高的情况会回退到更安全的路径。第三，它对代码块和表格做了专门处理，所以关键结构能够被保留下来。第四，它和真实产品行为是连接起来的，包括数据拉取、清洗、渲染、汇总以及兜底逻辑。用一句话总结，这个项目把经典的 LCS 思路变成了一个可用的富文本 diff 功能。"
    )

    add_heading(document, "准备小贴士", level=1)
    add_bullets(
        document,
        [
            "如果你担心紧张，可以把每页第一句单独记下来，作为开场锚点。",
            "如果 tutor 追问实现细节，可以重点补充 block-level LCS、long-text fallback、code block line diff、table diff 这四个点。",
            "如果 tutor 问项目价值，可以直接强调 readability、stability、specialization、product-readiness 这四个关键词。",
        ]
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
