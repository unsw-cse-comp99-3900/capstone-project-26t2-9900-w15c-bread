from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Git\capstone-project-26t2-9900-w15c-bread")
OUTPUT_DIR = ROOT / "outputs"
TECH_PATH = OUTPUT_DIR / "text-diff-technical-guide-latest-cn-2026-07-22.docx"
QA_PATH = OUTPUT_DIR / "tutor-qa-presentation-bilingual-2026-07-22.docx"

YELLOW = "FFD500"
BLACK = "111111"
DARK = "252525"
GREY = "666666"
LIGHT_GREY = "F2F3F5"
PALE_YELLOW = "FFF7CC"
RED = "B42318"
GREEN = "137333"
BLUE = "1F5AA6"
WHITE = "FFFFFF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.13

    for name, size, before, after in (
        ("Title", 26, 0, 10),
        ("Heading 1", 17, 14, 7),
        ("Heading 2", 13.5, 10, 5),
        ("Heading 3", 11.5, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = short_title
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.runs[0]
    hr.font.name = "Arial"
    hr.font.size = Pt(8.5)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor.from_string(GREY)
    p_pr = hp._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), YELLOW)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc, title, subtitle, meta_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("UNSW  ·  COMP9900  ·  TEAM W15C BREAD")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GREY)

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    r = p.add_run(title)
    r.font.color.rgb = RGBColor.from_string(BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(GREY)

    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    bar.autofit = False
    set_table_widths(bar, [0.18, 6.55])
    set_cell_fill(bar.cell(0, 0), YELLOW)
    set_cell_fill(bar.cell(0, 1), BLACK)
    bar.cell(0, 1).text = "LATEST CODE REVIEW · 22 JULY 2026"
    for cell in bar.rows[0].cells:
        set_cell_margins(cell, 100, 120, 100, 120)
    run = bar.cell(0, 1).paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph()
    for line in meta_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_page_break()


def add_callout(doc, title, body, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    set_table_widths(table, [6.65])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3.5)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3.5)
        p.add_run(item)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [6.65])
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F7F7F7")
    set_cell_margins(cell, 90, 120, 90, 120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)


def add_table(doc, headers, rows, widths=None, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        set_table_widths(table, widths)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, BLACK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(font_size)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            cell = cells[idx]
            if row_index % 2:
                set_cell_fill(cell, "FAFAFA")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    return table


def add_repo_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Code reference: ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GREY)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.2)
    r.font.color.rgb = RGBColor.from_string(BLUE)


def build_technical_doc():
    doc = Document()
    configure_document(doc, "Dynamic History · Text Diff Technical Guide")
    add_cover(
        doc,
        "Dynamic History — Text Diff 全面技术说明",
        "基于重新 pull 后最新版代码的实现级审阅（含算法、显示、恢复、安全与测试）",
        [
            "Repository: capstone-project-26t2-9900-w15c-bread",
            "Reviewed commit: 2e3e05718c37968df25642b68c2c82db39ebfde4",
            "Frontend verification: 11 test suites / 169 tests passed",
            "Scope: Confluence rich-text comparison and selective recovery end-to-end path",
        ],
    )

    doc.add_heading("0. 阅读结论（先看这一页）", level=1)
    add_callout(
        doc,
        "一句话定义",
        "Dynamic History 不是把 Confluence 页面降成纯文本后做 diff；它保留原始 Confluence Storage，另建安全的可读 HTML 预览，把页面拆成语义块，用 LCS 对齐不变块，再把变化投影成可选择的恢复单元。",
    )
    add_bullets(doc, [
        "比较方向固定为：Selected Historical → Current。removed 属于历史侧；added 属于当前侧。",
        "主页面级 diff 的正式结果类型只有 same / removed / added。显示层可把相邻 removed + added 配成一个 modified decision。",
        "恢复默认选择 Current；用户只有显式选择 old，才会把历史块放入虚拟 Draft。点击选择不会写页面。",
        "恢复粒度主要是 block。表格可以 cell-level 高亮，但写回仍以完整 table block 为选择单位。",
        "可读预览和可写回 Storage 严格分离；显示 HTML 从来不是恢复的数据源。",
        "发布前后端重新读取 Current version，并用 expectedVersionNumber 阻止覆盖并发更新。",
        "过大输入不会硬算矩阵：页面块 LCS 120,000 cells、inline 80,000、coarse sentence 240,000、code line 500,000 均有上限。",
    ])

    doc.add_heading("1. 审阅范围与代码地图", level=1)
    doc.add_paragraph(
        "本次按最新版仓库检查了 Forge manifest、后端 resolver、前端入口、两种比较视图、共享显示模型、Storage 重建、恢复工作流、样式和全部测试。与 text diff 直接相关的核心文件如下。"
    )
    add_table(doc, ["层", "文件", "职责"], [
        ("输入/后端", "src/index.js", "读取历史版本与 Current Storage；附件、用户；写回前并发版本校验。"),
        ("核心渲染与 diff", "static/confluence-dynamic-history/src/utils.js", "Storage → 安全预览；语义块；签名；LCS；空行、表格、布局和代码特殊逻辑。"),
        ("显示模型", "static/confluence-dynamic-history/src/diffDisplay.js", "变化运行配对、choice key、modified 显示分类、布局嵌套。"),
        ("Inline 视图", "components/ComparisonPanel.js", "生成 richDiff、展示 Summary、连接恢复工作流。"),
        ("Side-by-side", "components/splitDiffModel.js / SideBySideDiffView.js", "完整文档双栏投影、对齐、占位、方向选择。"),
        ("恢复", "recoveryStorage.js / useRecoveryWorkflow.js", "从 choice 重建原始 Storage、预览、确认写回。"),
        ("验证", "*.test.js", "11 suites / 169 tests，覆盖 renderer、diff、恢复、两种视图和边界。"),
    ], widths=[1.0, 2.5, 3.15])

    doc.add_heading("2. 端到端数据流", level=1)
    add_numbered(doc, [
        "Forge resolver 使用 asUser() 读取页面版本、Current 页面 Storage、附件地址和用户信息。",
        "前端以 selected historical Storage 为 oldHtml，以 Current Storage 为 currentHtml 调用 buildRichTextDiffHtml。",
        "prepareConfluenceHtml 只为可读展示转换 Storage/ADF，并执行 tag、attribute、style、URL allowlist 清洗。",
        "extractDiffBlocks 同时保留 prepared renderedHtml 与原始 raw/storage HTML，生成语义块和 canonical key。",
        "页面块序列用 LCS 对齐；正式输出 same / removed / added。",
        "diffDisplay 将连续变化块配对、分组，创建稳定 choice keys，并投影成 Inline 或 Side-by-side rows。",
        "用户选择 old/current；默认 current。Preview Draft 同时生成 readable preview 和 recoverable Storage。",
        "确认发布时 resolver 再读 live page，验证版本号，再以新版本号 PUT 回 Current page。",
    ])
    add_code(doc, "Historical Storage ─┐\n                    ├─ semantic blocks ─ LCS ─ diff blocks ─ display rows ─ choices\nCurrent Storage ────┘                                      │\n        original Storage fragments remain attached ────────┴─ recovery Storage ─ version-checked PUT")
    add_repo_ref(doc, "ComparisonPanel.js:551-657; utils.js:6026-6115; useRecoveryWorkflow.js:240-305; src/index.js:555-647")

    doc.add_heading("3. 为什么先做语义块，而不是直接比较 HTML 字符串", level=1)
    doc.add_paragraph(
        "Confluence Storage 的序列化细节会变化：attribute 顺序、临时 local-id/macro-id、等价标签（b/strong、i/em）、自闭合节点写法都可能不同。如果直接比较字符串，会把无意义的序列化差异误报为用户内容变化。"
    )
    add_table(doc, ["块/结构", "比较单元", "恢复说明"], [
        ("Paragraph / Heading / Blockquote", "完整块", "文本或格式改变通常表现为 old removed + new added。"),
        ("Ordinary list", "完整 list block", "避免拆散编号、嵌套和 start/value 语义。"),
        ("Task / Decision", "item-level", "恢复时重新组合成合法 task-list / decision-list Storage。"),
        ("Table", "完整 table block", "可靠时 cell-level 可视化；选择仍是整表。"),
        ("Panel / Macro / Image / Unsupported", "原子 raw block", "不允许从 fallback preview 反向重建。"),
        ("Compatible layout", "边界 + child blocks", "子内容可独立选择；列宽是独立原子选择。"),
        ("Incompatible layout", "完整 layout", "结构不一致时保守回退，避免破坏 wrapper。"),
    ], widths=[1.55, 1.45, 3.65])
    add_repo_ref(doc, "utils.js:3489-3569, 3662-3685, 3803-3975, 4147-4192")

    doc.add_heading("4. 可读预览与原始 Storage 的双轨模型", level=1)
    add_callout(
        doc,
        "核心安全设计",
        "previewHtml 用于人读；rawHtml / oldRawHtml / newRawHtml 用于写回。任何 unsupported fallback card、diff decoration、红绿标记都不能成为 Storage 数据源。",
        fill="EAF2F8",
    )
    add_bullets(doc, [
        "prepareConfluenceHtml 递归隔离 layout cells，避免正则跨 cell 吞掉后续内容。",
        "支持 Storage/ADF 中的 paragraph、heading、list、task、decision、table、panel、code、date、status、mention、emoji、image、expand、layout、whiteboard 等预览形态。",
        "sanitizer 只保留明确 tag 集；未知 tag 解包而不是执行。class 被移除，只提取允许的颜色/对齐提示。",
        "href/src 拒绝 javascript:；相对 URL 可绑定 baseUrl；link 强制 target=_blank、rel=noreferrer。",
        "unsupported block 显示 readable fallback + raw inspector，但恢复必须仍持有 raw Storage，否则写回被禁用。",
    ])
    add_repo_ref(doc, "utils.js:2850-3207; recoveryStorage.js:9-14, 112-127")

    doc.add_heading("5. Canonical signature：怎样判断两个 rich blocks 真正相同", level=1)
    add_bullets(doc, [
        "文本统一 NBSP/空白；code block 保留换行并 trimEnd，不把代码压成一行。",
        "b → strong、i → em；style declaration 排序，attribute 顺序不影响签名。",
        "class、aria-hidden、raw inspector、task marker 不参与用户内容身份。",
        "保留 href、src、alt、title、datetime、width/height、rowspan/colspan、safe style 及 data-dh-* 语义属性。",
        "Storage signature 忽略 volatile local-id、macro-id、node id；但保留真实 macro parameter、代码 CDATA 和图像元数据。",
        "unsupported block 的 key 基于 stable raw Storage hash，避免丢失真实参数变化。",
    ])
    add_repo_ref(doc, "utils.js:3210-3569")

    doc.add_heading("6. 页面级 LCS 算法", level=1)
    doc.add_paragraph(
        "buildRichTextDiffHtml 分别提取 oldBlocks 与 currentBlocks，并以 block.key 是否相等作为 LCS 匹配条件。DP 矩阵从右下向左上填充；回溯时，相等输出 same，否则根据未来 LCS 长度输出 removed 或 added。"
    )
    add_code(doc, "dp[i][j] = old[i].key === current[j].key\n  ? dp[i+1][j+1] + 1\n  : max(dp[i+1][j], dp[i][j+1])\n\nTie-break: dp[i+1][j] >= dp[i][j+1] ⇒ emit removed first")
    add_table(doc, ["输入情况", "输出"], [
        ("两侧都空", "空 diff。"),
        ("历史为空", "Current 全部 added。"),
        ("Current 为空", "Historical 全部 removed。"),
        ("oldCount × currentCount > 120,000", "不分配 DP；Current blocks 作为 same 预览，limited=true。"),
        ("一般情况", "same / removed / added 序列，再做 table decoration、blank-line 和 list-break compact。"),
    ], widths=[2.3, 4.35])
    add_callout(doc, "复杂度", "时间与空间均为 O(n×m)，所以必须有 120,000-cell 安全阈值。这里 n/m 是语义块数量，不是字符数量。")
    add_repo_ref(doc, "utils.js:6026-6115")

    doc.add_heading("7. “modified” 到底在哪里产生", level=1)
    doc.add_paragraph(
        "这是答辩最重要的精确表述。当前主页面 pipeline 只直接发出 same、removed、added。一个替换通常是相邻 old removed + new added。diffDisplay 再根据 nodeType、tag、空/非空兼容性和顺序，将相关两项共用一个 choice key，并把显示行标成 changeKind=modified。"
    )
    add_bullets(doc, [
        "可配对前提：removed/added、nodeType 相同、tag 相同、空块属性一致。",
        "连续 change run 内用 DP 配对；pair 的基础分是 100，再加 character-bigram similarity，保证尽量多配对并保序。",
        "未配对的 spacer blank blocks 会附到最近配对组，避免恢复文字后遗留/丢失关联空行。",
        "choice key 是原 diff block indices 拼接，如 12:13；所有块索引映射回相同 key，保证显示与写回一致。",
        "代码仍保留 direct modified helpers（如 line/code/table），且共享显示层兼容 direct modified；但不要误称主页面 LCS 普遍直接生成 modified block。",
    ])
    add_repo_ref(doc, "diffDisplay.js:3-108, 150-220, 293-328; utils.js:5896-5975")

    doc.add_heading("8. Inline、Side-by-side 与统计", level=1)
    add_table(doc, ["视图", "数据源", "表现"], [
        ("Inline Summary", "diffDisplay.rows", "same 正常显示；变化显示 -/+；点击后选择 Keep current / Restore old / Undo。"),
        ("Side-by-side", "同一个 display → split rows", "Historical 左、Current 右；unchanged 两侧均显示；单侧变化用 placeholder 对齐。"),
        ("Version Difference Notes", "Current → virtual Draft 再跑 diff", "回答最终 Draft 相对 Current 真正会改变什么。"),
        ("Comments summary", "canonical summary", "added/removed 来自 unit count；modifiedBlocks 来自 display selectable rows。"),
    ], widths=[1.35, 1.75, 3.55])
    add_bullets(doc, [
        "Side-by-side 每个 visual row 同时拥有左右 pane，较高一侧决定共享 row height，后续块继续对齐。",
        "Historical-only：左侧内容 + 右侧 Not present；Current-only 反之。",
        "简单单元素 modified row 可再做 side-specific inline highlight；若 HTML 有复杂子节点，则保守保持完整块。",
        "分隔线可拖动，30%–70% clamp；键盘可调；双击回 50%。",
        "两个视图共用 recoveryChoices，因此切换视图不会改变同一版本对的选择语义。",
    ])
    add_repo_ref(doc, "splitDiffModel.js:11-257; SideBySideDiffView.js:24-35, 75-117, 448-531; App.js:125-150, 255-296")

    doc.add_heading("9. 特殊内容的 diff 细节", level=1)
    doc.add_heading("9.1 空行与 Enter/Shift+Enter", level=2)
    add_bullets(doc, [
        "连续空 editor paragraphs 折叠为 blank_line_run，key 包含 layoutPath 与 count。",
        "识别 p、top-level br、ADF hardBreak/content、空 formatting span、caret Unicode controls；不吞 macro/media/link/mention。",
        "2 → 5 只报告 3 blank lines added；5 → 2 只报告 3 removed。块仍保留两侧完整 run Storage，恢复精确数量。",
        "非空 paragraph 内的 br（常见 Shift+Enter）仍属于该 paragraph，不当作独立 Enter 空行。",
        "若只是 list 之间的空行数量造成 list 被拆/合，list_break_change 会把相同 list sequence 与空行 delta 合成一个变化。",
    ])
    add_repo_ref(doc, "utils.js:3978-4144, 4385-4688")

    doc.add_heading("9.2 表格", level=2)
    add_bullets(doc, [
        "用逻辑 grid coordinate 匹配 cell；rowspan/colspan 会占据所有 logical slots。嵌套 table rows 被排除。",
        "same geometry 可在一个表中展开 changed cell 的 previous/current 区域，并分别保留原背景色。",
        "只支持保守的结构变化：末尾追加/删除完整行、最右追加/删除完整列，以及可验证的 terminal row+column L-shape。",
        "中间插入、稳定轴位移、span 跨界、重复 coordinate、ambiguous geometry 会 whole-table fallback。",
        "cell-level 是 display fidelity，不是独立 staging；Keep/Restore 仍选择 complete old/current table Storage。",
    ])
    add_repo_ref(doc, "utils.js:5047-5893; recoveryDiffDisplay.js:65-151")

    doc.add_heading("9.3 Layout、Task、Decision、Code 与 unsupported", level=2)
    add_bullets(doc, [
        "只有 old/current layout structure signature 相同才拆 layout；边界本身 non-selectable，child blocks 可选择。",
        "column width change 是独立 choice。恢复 old width 时从 Current opening tag 仅替换 width attribute，保留 Current local IDs。",
        "Task/Decision item 独立选择，但 Storage 重建会重新组成一个合法 group；完整同源 Decision group 可原样复用。",
        "Code macro 的显示可 line diff；写回前规范 CDATA，已合法 CDATA 尽量保持原样。",
        "Unsupported/raw block 原子化；没有 raw Storage 就返回 blocking error，宁可不写也不丢数据。",
    ])
    add_repo_ref(doc, "utils.js:2737-2850, 3803-3975; recoveryStorage.js:34-107, 180-680")

    doc.add_heading("10. 恢复选择、Preview Draft 与写回", level=1)
    add_bullets(doc, [
        "default = current：没有显式 choice 的 row 始终保留 live Current 内容。",
        "added + old ⇒ omitted；removed + current ⇒ omitted；blank_line_count_change 两侧都真实存在，不能按普通单侧块省略。",
        "Preview 同时保存 selected/current version number、每个 row 的 block indices 与 choice、previewHtml、storageHtml、storageError。",
        "切换版本 pair 会 reset choices/preview/write status，并忽略旧比较的 stale callback。",
        "用户第一次点击只改变虚拟 Draft；打开 Preview 仍是只读；第二次明确点击 Write to Current Page 才 invoke resolver。",
        "resolver 验证 pageId、非空字符串、2 MB limit；asUser() 读取 live page；expectedVersionNumber 必须匹配；再 PUT v+1。",
    ])
    add_repo_ref(doc, "recoveryStorage.js:16-107, 637-680; useRecoveryWorkflow.js:12-86, 139-160, 212-342; src/index.js:555-687")

    doc.add_heading("11. 安全阈值、降级与错误状态", level=1)
    add_table(doc, ["位置", "阈值", "行为"], [
        ("Page block LCS", "120,000 cells", "Current-side safe preview；limited=true；不提供不可靠的 per-block merge。"),
        ("Inline token LCS", "80,000 cells", "转 coarse sentence diff；仍过大则 limited。"),
        ("Coarse sentence LCS", "240,000 cells", "返回 Current text，无伪精细 highlight。"),
        ("Side split inline", "40,000 cells", "不做额外 side word highlight。"),
        ("Code line LCS", "500,000 cells", "Current lines 作为 safe same；limited=true。"),
        ("Write payload", "2,000,000 chars", "resolver 拒绝。"),
        ("Mention resolution", "最多 100 account IDs", "超出不阻止 diff；显示 safe fallback。"),
    ], widths=[1.6, 1.35, 3.7])
    add_callout(doc, "需要修正文案", "Inline UI 当前 limited warning 写成“safer line-level comparison”，但页面级 >120,000 的实际行为是 Current-side safe result，而不是可靠的 full line-level merge。答辩应按代码解释，并建议后续改文案。", fill="FDECEC")

    doc.add_heading("12. 测试证据", level=1)
    doc.add_paragraph("2026-07-22 在 frontend 目录以 react-scripts 运行完整测试：")
    add_code(doc, "Test Suites: 11 passed, 11 total\nTests:       169 passed, 169 total\nSnapshots:   0 total\nTime:        28.061 s")
    add_bullets(doc, [
        "utils.test.js：renderer、semantic block、signature、layout、blank lines、tables、ADF/Storage、security-related normalization。",
        "recoveryStorage.test.js：task/decision group、layout width、unsupported raw block、CDATA、duplicate prevention。",
        "diffDisplay/splitDiffModel：pairing、indices、full-document alignment、side-specific HTML。",
        "ComparisonPanel/SideBySide：共享选择、preview/write workflow、limited/error、responsive/accessibility。",
        "仍缺后端 resolver 自动化测试；writeRecoveredPage 目前主要靠代码级防护和前端契约测试。",
    ])
    add_callout(doc, "Dependency note", "package.json 没有 test script；本次使用 node_modules/.bin/react-scripts.cmd。npm ci 需要 --legacy-peer-deps，且安装时出现 Create React App/Babel 维护警告。它不影响本次 169 tests 通过，但属于技术债。", fill=LIGHT_GREY)

    doc.add_heading("13. 当前限制与可改进点", level=1)
    add_table(doc, ["限制", "影响", "建议"], [
        ("O(n×m) DP 矩阵", "大页必须 limited。", "改用 Myers/Hirschberg/patience anchors，降低空间或提高可比较规模。"),
        ("恢复主要 block-level", "不能只恢复复杂 paragraph 的一个词或 table 的一个 cell。", "在确保 Storage validity 后引入受限的 sub-block staging。"),
        ("Manual renderer", "必须持续追踪 Confluence 新 Storage/ADF 类型。", "建立真实页面 corpus、contract tests 与 telemetry；评估官方 conversion API。"),
        ("Table conservative fallback", "中间插行/列会整表展示。", "研究 row/column key matching，但仍需可解释的 ambiguity rules。"),
        ("前端 DOMParser + dangerouslySetInnerHTML", "依赖 sanitizer 完整性。", "持续 fuzz/XSS tests；考虑成熟 sanitizer 或 Trusted Types。"),
        ("后端测试空白", "并发/REST error 边界缺少自动回归。", "增加 resolver unit tests 和 Forge integration tests。"),
    ], widths=[1.55, 1.7, 3.4])

    doc.add_heading("14. PPT 表述与最新版代码的一致性", level=1)
    add_table(doc, ["PPT 内容", "代码结论", "答辩表述"], [
        ("Block-based + LCS", "准确。", "补充 canonical signature、120k threshold 与 same/removed/added contract。"),
        ("Restore one paragraph", "准确，但以 block 为恢复单位。", "不要暗示任意字符级或 table cell 级写回。"),
        ("Preview vs source data", "准确且是核心安全设计。", "说明 preview HTML never becomes Storage。"),
        ("Check latest version", "准确。", "expectedVersionNumber + asUser() live GET + PUT v+1。"),
        ("80+ considered / 34 priority", "属于研究与测试范围陈述。", "不要说 runtime 自动支持 80 或精确保证 34；准备研究表作为证据。"),
        ("Bulk actions", "准确。", "App 共享 choice state；choose-all/reset/preview 在 workspace toolbar。"),
        ("7-slide section", "实际文件 9 页：title + 7 content + Q&A。", "可以说 seven product-design content slides。"),
        ("Demo handover", "Slide 8 仍有 Teammate placeholders。", "正式演示前替换姓名并确认顺序。"),
    ], widths=[1.75, 2.05, 2.85])

    doc.add_heading("附录 A：答辩时可背的 30 秒算法解释", level=1)
    add_callout(doc, "English answer", "We first convert each Confluence version into semantic blocks while keeping the original Storage fragments. We build canonical signatures that ignore serialization noise but retain meaningful formatting and metadata. Then an LCS aligns unchanged blocks, and unmatched blocks become removed or added. The UI may pair a related removed and added block as one modified decision, but recovery still selects complete original Storage blocks. This keeps the comparison readable and the write-back safe.")
    doc.add_paragraph("中文记忆：语义块 → 稳定签名 → LCS 对齐 → 显示层配对 → 原始 Storage 恢复 → 版本校验写回。")

    doc.add_heading("附录 B：关键代码索引", level=1)
    add_table(doc, ["主题", "位置"], [
        ("sanitized preview renderer", "utils.js:2850-3207"),
        ("canonical signature", "utils.js:3210-3569"),
        ("block extraction / blank lines", "utils.js:3803-4192"),
        ("page-level LCS", "utils.js:6026-6115"),
        ("line LCS", "utils.js:6138-6203"),
        ("table compatibility", "utils.js:5047-5893"),
        ("display pairing / choice keys", "diffDisplay.js:13-328"),
        ("split projection", "splitDiffModel.js:11-257"),
        ("Storage reconstruction", "recoveryStorage.js:16-680"),
        ("preview/write workflow", "useRecoveryWorkflow.js:139-342"),
        ("backend concurrency check", "src/index.js:555-687"),
    ], widths=[2.4, 4.25])

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(TECH_PATH)


QA_ITEMS = [
    ("Problem & value", "1. What problem are you solving?", "Confluence can compare and restore versions, but recovering only one useful part is still cumbersome. Users often restore the whole page or copy content manually, which can overwrite or miss newer work. Dynamic History turns version review into a selective recovery workflow.", "先说 selective recovery gap，再说全量恢复/手工复制的风险。对应 PPT slide 2。"),
    ("Problem & value", "2. How is this different from Confluence's native version history?", "Native history is our source of versions. Our contribution is the review-and-recovery layer: semantic rich-text comparison, block-level choices, a combined preview, and a version-checked write-back as a new page version.", "不要说替代 Confluence；强调建立在 native history 之上。"),
    ("Problem & value", "3. Why is selective recovery valuable?", "Most users do not want to roll back an entire page. They want one deleted paragraph, table, decision, or layout change while retaining newer edits elsewhere. Selective recovery reduces the cost and risk of that task.", "用真实场景：恢复误删的一段，同时保留之后新增内容。"),
    ("Evolution", "4. What changed from Sprint 1 to Sprint 2?", "Sprint 1 established the version timeline, version selection, and basic comparison. Sprint 2 expanded rich-text handling and added the complete workflow: clearer review, shared choices, preview, comments, selective recovery, and safe publishing.", "对应 slide 3；不要只列功能，要说 structure → workflow。"),
    ("Research", "5. What do the 80+ and 34 numbers mean?", "They describe our design research scope: we considered more than 80 content types and prioritised 34 for deeper testing. They are not a claim that the runtime automatically guarantees every one of those types. The implementation also keeps unsupported content as safe raw blocks.", "这是高风险数字题。若 tutor 要证据，应展示研究清单；不要把数字说成代码 coverage 计数。"),
    ("Algorithm", "6. Why did you choose block-based comparison?", "A Confluence page contains meaningful structures, not just characters. Treating a paragraph, list, table, panel, or image as a block gives users understandable recovery units and avoids breaking rich-text structure during write-back.", "关键词：meaningful unit、valid Storage、explainable choice。"),
    ("Algorithm", "7. How does the LCS algorithm work here?", "We create a sequence of canonical block signatures for each version. LCS finds the longest ordered set of equal signatures. Matched blocks are unchanged; unmatched historical blocks are removed and unmatched current blocks are added.", "20 秒版本；若追问复杂度，再答下一题。"),
    ("Algorithm", "8. What is the complexity of your diff?", "The page-level dynamic program is O(n times m) in time and space, where n and m are block counts. We cap the matrix at 120,000 cells. Large inputs use a limited safe current-side result instead of risking excessive memory.", "明确阈值 120,000；不说无限可扩展。"),
    ("Algorithm", "9. Does the main algorithm output a modified block?", "The active page-level contract outputs same, removed, and added. A replacement is normally an old removed block followed by a new added block. The display model can pair those two into one modified decision with one recovery key.", "这是最能体现真正读过代码的回答。"),
    ("Algorithm", "10. How do you avoid false differences caused by HTML serialization?", "We use canonical signatures. They normalise whitespace, treat equivalent tags such as b and strong consistently, sort style and attribute information, and ignore volatile Confluence IDs. Meaningful links, dates, image metadata, formatting, spans, and macro parameters remain part of the identity.", "强调 ignore noise, retain semantics。"),
    ("Algorithm", "11. What if only formatting changes but text stays the same?", "Formatting can still change the canonical DOM or Storage signature, so the containing block is treated as changed. Recovery selects the complete original block, which preserves the chosen version's formatting.", "不是纯 text diff；rich semantics 会参与。"),
    ("Algorithm", "12. How do you compare Chinese text?", "Inline tokenisation treats individual CJK characters as units, while Latin letters and numbers are grouped into word-like tokens. At the page level, however, alignment is still based on canonical semantic block signatures.", "区分 page-level block LCS 与辅助 inline tokenisation。"),
    ("Granularity", "13. Can users restore one word inside a paragraph?", "Not in the current recovery model. The product can visually highlight smaller differences in some simple cases, but the safe recovery unit is the complete block. This protects Confluence structure and makes the action predictable.", "诚实边界：显示细，恢复粗。"),
    ("Granularity", "14. Are table cells independently recoverable?", "No. Compatible tables can show cell-level differences, including rowspan and colspan-aware matching, but Keep or Restore still selects the complete old or current table block. Cell-level staging is future work.", "必答题：cell-level display ≠ cell-level write-back。"),
    ("Tables", "15. How do you compare tables safely?", "We build a logical grid, so merged cells occupy their rowspan and colspan coordinates. We support same-shape changes and conservative terminal row or column changes. Ambiguous middle insertions or incompatible spans fall back to complete tables.", "突出 conservative fallback。"),
    ("Special cases", "16. Why do blank lines need special handling?", "Several Confluence Storage forms can represent the same editor blank line. We collapse consecutive blank paragraphs into a count-aware run. A change from two to five blank lines becomes three added, while both complete runs remain available for exact recovery.", "可用 2→5 的例子。"),
    ("Special cases", "17. How do lists, tasks and decisions stay valid after mixed recovery?", "Task and Decision items can be selected individually, but write-back does not concatenate loose item HTML. The recovery layer reconnects related groups and rebuilds a valid task-list or decision-list wrapper, reusing the original complete group when possible.", "解释显示粒度与 Storage 合法性之间的桥梁。"),
    ("Layouts", "18. How do you handle Confluence multi-column layouts?", "If the historical and current layout structures match, we keep non-selectable layout boundaries and diff child blocks independently. Column widths are a separate atomic choice. If the structure differs, we fall back to the whole layout.", "可提 25/75、25/50/25 width preservation。"),
    ("Safety", "19. Why separate the readable preview from original Confluence data?", "The preview is transformed and decorated for humans, so it is not a trustworthy write-back source. We keep original Storage fragments attached to every block and reconstruct from those fragments. This preserves formatting and avoids writing our red or green diff markup into Confluence.", "对应 slide 6 的核心。"),
    ("Safety", "20. Is dangerouslySetInnerHTML safe in your app?", "It is only used with HTML produced by our manual renderer and sanitizer. The sanitizer allowlists tags, attributes, styles, colours and URLs, removes classes, and rejects javascript URLs. Unsupported data is escaped and kept separately as raw Storage. We would still add fuzz and XSS regression testing as defence in depth.", "不要只说“safe”；要说具体控制 + 后续改进。"),
    ("Safety", "21. What happens when you do not understand a Confluence content type?", "We do not silently drop it. The UI shows a readable fallback and optional raw inspector, while the original Storage remains an atomic recovery block. If raw Storage is missing, write-back is disabled to avoid data loss.", "Fail safe, not fail open。"),
    ("Safety", "22. How do you prevent overwriting somebody else's newer edit?", "The preview records the current version number. On confirmation, the backend uses asUser to read the live page again. If the live version no longer matches the expected version, the write is rejected and the user must reload.", "对应 expectedVersionNumber；强调 optimistic concurrency guard。"),
    ("Safety", "23. What permissions are used?", "The backend calls Confluence asUser, so Confluence applies the invoking user's permissions. The manifest uses the page, attachment, user-read and page-write scopes required by the current workflow. We avoid asApp for user-context page operations.", "可引用 read:page, read:attachment, read:confluence-user, write:page。"),
    ("Workflow", "24. What happens when a user clicks Restore?", "It only changes a choice in a virtual Draft. The source panes remain read-only. The user then opens Preview Draft, reviews the reconstructed result and its difference from Current, and must explicitly confirm Write to Current Page.", "选择不是写操作；two-step confirmation。"),
    ("Workflow", "25. What is the default if the user makes no choice?", "Current content wins by default. Historical content is used only after an explicit old choice. This makes the workflow conservative and protects newer work.", "一句话说清 default current。"),
    ("Workflow", "26. Do Inline and Side-by-side produce different recovery results?", "No. They share the same canonical display rows, choice keys and recovery state for one version pair. They are different visual projections of the same comparison and write-back model.", "共享 diffDisplay + recoveryChoices。"),
    ("UX", "27. How did client feedback change the design?", "We reduced repeated checkbox-style actions, introduced clearer shared controls and bulk actions, and changed change styling to borders and labels so original background colours remain visible. The goal was lower selection cost and unambiguous states.", "对应 slide 7；可说 feedback → design implication。"),
    ("UX", "28. Why provide both Inline and Side-by-side views?", "Inline is efficient for scanning changes in one flow. Side-by-side provides complete historical and current context, keeps corresponding blocks aligned, and makes source ownership clearer. Because they share the model, users can choose the reading style without changing recovery semantics.", "不要说冗余；强调不同 review task。"),
    ("Testing", "29. How have you tested the implementation?", "The latest frontend suite passes 11 test suites and 169 tests. Coverage includes semantic rendering, signatures, blank lines, tables, layouts, tasks, Decisions, raw unsupported blocks, recovery Storage, both views, accessibility and error or limited states.", "给真实数字 11/169；再承认 resolver tests 是缺口。"),
    ("Testing", "30. What important testing is still missing?", "The largest gap is automated backend resolver and live Forge integration testing. We also want a larger corpus of real Confluence Storage, cross-browser checks, performance benchmarks, and security fuzz tests for the sanitizer.", "主动说缺口会比声称“fully tested”更可信。"),
    ("Limits", "31. What is the biggest current technical limitation?", "The safest recovery unit is still the block, and the page-level LCS uses an O(n times m) matrix. That limits very large pages and prevents fine-grained recovery inside complex blocks. Both are deliberate safety trade-offs for Sprint 2.", "trade-off，而不是隐藏缺陷。"),
    ("Limits", "32. What would you improve next?", "I would first add backend and live integration tests, then benchmark a lower-memory sequence diff such as Myers or Hirschberg. After that, I would explore safe sub-block recovery for selected paragraphs or table cells, backed by Storage validity checks.", "优先级：safety/testing → performance → finer recovery。"),
    ("Product", "33. How would you measure whether the product succeeds?", "We would measure time to recover a selected change, the number of manual copy steps avoided, preview-to-publish completion, recovery errors, aborted writes caused by concurrent edits, and user confidence in identifying the correct change.", "从功能 demo 转向 measurable outcomes。"),
    ("Demo", "34. What should we show in the live demo?", "Use one page containing a paragraph edit, a table or rich-text change, and newer Current content. Show Inline, Side-by-side, Restore one historical block, Preview Draft, Version Difference Notes, and the final version-checked publish. Also prepare an unsupported or large-content explanation as a fallback.", "demo 应证明 selective + safe，不要只展示 UI。"),
    ("Demo", "35. What if the demo renderer fails?", "The comparison is wrapped in an error boundary and distinguishes renderer errors, limited results and empty diffs. We should explain the state, avoid write-back, reload with a known test page, and continue with screenshots rather than pretending the comparison is valid.", "准备 graceful fallback。"),
    ("Presentation", "36. Is the deck really seven slides?", "It contains nine slides in total: a title slide, seven product-design content slides, and a final questions slide. So the accurate phrase is 'a seven-slide product-design section' rather than a seven-slide file.", "用于避免 tutor 发现页数不一致。"),
]


def build_qa_doc():
    doc = Document()
    configure_document(doc, "Dynamic History · Tutor Q&A")
    add_cover(
        doc,
        "Dynamic History — Tutor Q&A 答辩准备",
        "根据最新版代码与 Dynamic_History_UNSW_Product_Design_revised.pptx 整理",
        [
            "Format: probable question + concise English answer + Chinese coaching note",
            "Deck reviewed: 9 slides including speaker notes",
            "Code verified: 11 test suites / 169 tests passed",
            "Recommended answer length: 20–40 seconds unless the tutor asks a follow-up",
        ],
    )

    doc.add_heading("使用方法", level=1)
    add_bullets(doc, [
        "先背每题英文第一句；它应该直接回答问题。第二、三句用于补证据。",
        "不要把“可视化粒度”说成“恢复粒度”：table cell 可高亮，但当前写回是 whole table block。",
        "不要把 80+ / 34 说成 runtime guarantee；它们是 research and prioritisation scope。",
        "遇到 limitation 题，用 deliberate safety trade-off 回答，再给 next step。",
        "所有安全题都回到三点：sanitized preview、original Storage、version-checked write-back。",
    ])
    add_callout(doc, "万能结构", "Answer → Evidence from implementation → Trade-off or limitation → Next step（如果被追问）。")

    doc.add_heading("开场 45 秒版本", level=1)
    add_callout(
        doc,
        "Suggested opening",
        "Dynamic History helps Confluence users recover selected rich-text content without rolling back an entire page. We preserve the original Confluence Storage, create a separate readable comparison, align semantic blocks with LCS, and let the user choose historical or current blocks in a virtual Draft. Nothing is written until the user reviews the Draft and confirms publishing, and the backend checks that the live page version has not changed.",
    )

    current_section = None
    for section, question, answer, coaching in QA_ITEMS:
        if section != current_section:
            doc.add_heading(section, level=1)
            current_section = section
        doc.add_heading(question, level=2)
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        r = p.add_run("Suggested answer: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GREEN)
        p.add_run(answer)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.keep_together = True
        r = p.add_run("中文提示：")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(coaching)

    doc.add_heading("PPT 逐页潜在追问", level=1)
    add_table(doc, ["Slide", "Tutor 可能追问", "准备重点"], [
        ("1 Title", "Who owns which part? Why 2m40s?", "确认最终 speaker、handover 与时间；不要保留占位。"),
        ("2 Problem", "Why is native restore insufficient?", "Selective recovery gap + risk of losing newer work。"),
        ("3 Evolution", "What is demonstrably new in Sprint 2?", "rich renderer、shared choices、preview/write、comments、side-by-side。"),
        ("4 Research", "Where did 80+ and 34 come from?", "带 research matrix；说明 considered vs prioritised vs implemented。"),
        ("5 Block/LCS", "Why LCS? Complexity? Modified?", "ordered alignment、120k cap、display pairing。"),
        ("6 Safe Recovery", "How is data preserved? Race condition?", "original Storage、sanitizer、expectedVersionNumber。"),
        ("7 Feedback", "What evidence did client provide?", "准备反馈记录或会议纪要；说具体 design response。"),
        ("8 Demo", "What exactly will each teammate show?", "替换 Teammate 1/2/3；明确 demo ownership。"),
        ("9 Q&A", "Any limitations/future work?", "block-level recovery、large-page cap、backend test gap。"),
    ], widths=[0.9, 2.55, 3.2])

    doc.add_heading("不要这样回答", level=1)
    add_table(doc, ["不准确说法", "改成"], [
        ("We do text diff character by character.", "We align semantic rich-text blocks; some views add limited inline highlighting."),
        ("The algorithm outputs modified blocks.", "The active page contract outputs same/removed/added; the display layer pairs related old/new blocks as modified decisions."),
        ("Users can restore any table cell.", "Cells can be compared visually; recovery currently selects the complete table."),
        ("We support all 80+ content types.", "We considered 80+ and prioritised 34; unsupported types are preserved as raw atomic blocks."),
        ("The preview is the page we write back.", "The preview is for readability; original Storage fragments are reconstructed for write-back."),
        ("It can never overwrite changes.", "It prevents a stale write when the version number changes; broader collaboration conflicts still require reload and review."),
        ("All tests are complete.", "169 frontend tests pass; backend and live integration automation remain important next work."),
    ], widths=[2.65, 4.0])

    doc.add_heading("答辩前 10 分钟检查单", level=1)
    add_bullets(doc, [
        "Slide 8 替换 [Teammate 1/2/3] 与“Replace demo owner names…”占位说明。",
        "带上 80+/34 的 research matrix 或至少截图，避免数字无来源。",
        "准备一个稳定 demo page：paragraph + table/rich block + newer Current content。",
        "预先选好 historical/current version numbers，确认 comparison direction。",
        "演示 Restore 后一定打开 Preview Draft，再发布；不要直接跳过 safety story。",
        "准备离线截图/录像，以防 Forge/network/demo site 失败。",
        "团队统一术语：Historical / Current / virtual Draft / block / Storage / LCS。",
        "全员统一限制表述：block-level recovery; table cell-level display only; large-page limited mode。",
    ])

    doc.add_heading("最后 20 秒总结", level=1)
    add_callout(
        doc,
        "Suggested closing",
        "Our main design contribution is not only showing differences. It is connecting rich-text comparison to a conservative recovery workflow. The user can understand the change, select a meaningful block, preview the combined result, and publish only after the system validates the live version. That is how we protect both formatting and newer work.",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(QA_PATH)


if __name__ == "__main__":
    build_technical_doc()
    build_qa_doc()
    print(TECH_PATH)
    print(QA_PATH)
