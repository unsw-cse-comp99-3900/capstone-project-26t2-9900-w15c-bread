from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("outputs/text-diff-algorithm-summary.docx")


def set_run_font(run, size, bold=False, color=None):
    """Apply both Western and East Asian fonts so Chinese text renders cleanly."""
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, 22, bold=True, color=RGBColor(0x0B, 0x25, 0x45))


def add_subtitle(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=14, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, 11, bold=False, color=RGBColor(0x55, 0x55, 0x55))


def add_heading(document, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {
        1: RGBColor(0x2E, 0x74, 0xB5),
        2: RGBColor(0x2E, 0x74, 0xB5),
        3: RGBColor(0x1F, 0x4D, 0x78),
    }
    spacing_before = {1: 16, 2: 12, 3: 8}
    spacing_after = {1: 8, 2: 6, 3: 4}

    paragraph = document.add_paragraph()
    set_paragraph_spacing(
        paragraph,
        before=spacing_before[level],
        after=spacing_after[level],
        line=1.1,
    )
    run = paragraph.add_run(text)
    set_run_font(run, sizes[level], bold=True, color=colors[level])


def add_body(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, before=0, after=4, line=1.2)
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, before=0, after=4, line=1.2)
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=10.5, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def add_algorithm_table(document):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.1), Inches(1.75), Inches(1.45), Inches(2.2)]
    headers = ["层级", "核心函数", "策略", "输出/效果"]

    for index, width in enumerate(widths):
        table.columns[index].width = width

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        set_cell_text(
            cell,
            header,
            size=10.5,
            bold=True,
            color=RGBColor(0x1F, 0x4D, 0x78),
        )

    rows = [
        ["块级", "buildRichTextDiffHtml", "LCS 对齐块", "决定 same / added / removed / modified"],
        ["行内", "buildInlineTextDiff", "按 token 做 LCS", "给段落、标题、列表项打增删标记"],
        ["长文本回退", "buildCoarseTextDiff", "先按句/行比，再对子段做行内 diff", "避免浏览器卡顿"],
        ["代码块", "buildCodeBlockDiff + buildLineDiff", "按行比较并保留缩进", "适合 `<pre><code>` 宏"],
        ["表格", "buildTableDiff", "同形状做单表格 cell diff；异形状并排展示", "避免破坏表结构"],
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def add_code_refs_table(document):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)

    for index, header in enumerate(["文件", "作用"]):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, size=10.5, bold=True)

    rows = [
        ["README.md", "对整体 diff 流程、返回结构和近期增强做总说明。"],
        ["static/confluence-dynamic-history/src/utils.js", "算法主实现：预处理、块级比对、行内 diff、代码块 diff、表格 diff。"],
        ["static/confluence-dynamic-history/src/components/ComparisonPanel.js", "消费 summary，展示变更计数、limited 提示和兜底错误预览。"],
        ["src/index.js", "从 Confluence 拉取版本、附件和作者信息，为前端 diff 提供输入。"],
    ]

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0])
        set_cell_text(cells[1], row[1])


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("基于仓库代码整理 | Text Diff Brief")
    set_run_font(run, 9, color=RGBColor(0x66, 0x66, 0x66))


def build_document():
    document = Document()
    configure_page(document)
    add_footer(document.sections[0])

    add_title(document, "基于代码的文本差异算法与配套操作总结")
    add_subtitle(
        document,
        "项目：Dynamic History（Confluence 版本比较） | 目标：支撑 3 分钟以内的技术讲解与演示",
    )

    add_body(
        document,
        "这套实现不是单一的“文本 diff 函数”，而是一条面向 Confluence 富文本的分层处理链路。它先把 storage HTML 规范化，再按块对齐内容，随后根据块类型选择行内 diff、代码行 diff 或表格 diff，最后把结果打包成可直接渲染的 HTML 和结构化 summary。"
    )

    add_heading(document, "一、核心结论", level=1)
    add_bullets(
        document,
        [
            "算法主线是“预处理 -> 块级 LCS 对齐 -> 类型化细分 diff -> 汇总渲染”，而不是一次性全文本比对。",
            "实现重点不只是“算出差异”，还包括附件解析、Confluence 宏展开、安全渲染、性能阈值控制和前端兜底展示。",
            "设计目标偏工程实用：优先保证可读性、稳定性和浏览器可承受性，而不是追求最细粒度、最昂贵的最优 diff。",
        ],
    )

    add_heading(document, "二、算法主流程", level=1)
    add_numbered(
        document,
        [
            "用 prepareConfluenceHtml 对 Confluence storage HTML 做预处理：展开链接、代码宏、表情和图片附件，并过滤危险标签/属性。",
            "用 extractDiffBlocks 把页面拆成可比较块，块元数据由 extractBlockMeta 生成，包含 key、nodeType、text、html 等字段。",
            "在 buildRichTextDiffHtml 中对 oldBlocks 与 currentBlocks 做块级 LCS，比出 same / added / removed。",
            "若两个块不完全相同，但 canPairForInlineDiff 判断“适合进一步比较”，就转入 modified 路径。",
            "modified 路径再按类型分流：普通文本做 token diff，长文本启用 coarse fallback，代码块按行 diff，表格按形状决定 cell diff 或并排展示。",
            "最终由 buildDiffResult 统一生成 html、blocks、summary、added、removed、limited，供 UI 直接渲染。"
        ],
    )

    add_heading(document, "三、各类差异算法与阈值", level=1)
    add_algorithm_table(document)

    add_heading(document, "四、为什么它能在前端跑得住", level=1)
    add_bullets(
        document,
        [
            "块级 diff 限制：buildRichTextDiffHtml 里 oldCount * currentCount 超过 120000 时，会把当前内容按 same 输出并标记 limited。",
            "行内 diff 限制：buildInlineTextDiff 里 token DP 超过 80000 个单元时，优先退回 buildCoarseTextDiff。",
            "长文本 fallback：buildCoarseTextDiff 先按句子或换行切片，再对子片段做局部 inline diff；若仍太大，直接输出 safer preview。",
            "代码块限制：buildLineDiff 使用按行 LCS，并设置 500000 单元上限，避免大代码块拖垮渲染。",
        ],
    )

    add_heading(document, "五、配套操作（工程侧同样关键）", level=1)
    add_bullets(
        document,
        [
            "数据准备：src/index.js 负责分页拉取页面版本、附件下载地址和作者显示名，给前端提供可比较输入。",
            "富文本清洗：Confluence 链接、code 宏、image 宏先展开，避免前端直接面对原始 storage 格式。",
            "安全渲染：只保留允许标签和属性，屏蔽 javascript: 之类的危险值，并补全相对链接。",
            "结果汇总：ComparisonPanel.js 直接消费 summary，显示 additions、removals、modifiedBlocks、total changes。",
            "异常兜底：如果 diff 渲染异常，ComparisonPanel 会退回到安全提示，而不是整页崩溃。",
            "样式配合：styles.css 中的 diff block、inline、code line、table panel 样式让不同类型变更在视觉上可区分。",
        ],
    )

    add_heading(document, "六、代码位置速查", level=1)
    add_code_refs_table(document)

    add_heading(document, "七、3 分钟演讲建议", level=1)
    add_bullets(
        document,
        [
            "第 1 分钟：先讲整体架构，强调这是“分层 diff 流水线”，不是单函数。",
            "第 2 分钟：重点讲 4 个核心策略：块级 LCS、token 行内 diff、长文本 fallback、表格/代码特化。",
            "第 3 分钟：讲配套操作和工程价值，说明为什么它在富文本、附件、宏、性能限制下仍能稳定工作。",
        ],
    )

    add_body(
        document,
        "一句话总结：这套实现的亮点不在于发明了全新的 diff 理论，而在于把经典 LCS 思路与富文本预处理、类型化比对、性能保护和前端渲染契合地组合成了一套能上线的工程方案。"
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
