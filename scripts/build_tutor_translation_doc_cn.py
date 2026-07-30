from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("outputs/tutor-text-diff-translation-cn.docx")


def set_run_font(run, size, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.2):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, 22, bold=True, color=RGBColor(0x0B, 0x25, 0x45))


def add_subtitle(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=14, line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, 11, color=RGBColor(0x55, 0x55, 0x55))


def add_heading(document, text, level=1):
    sizes = {1: 16, 2: 13}
    colors = {1: RGBColor(0x2E, 0x74, 0xB5), 2: RGBColor(0x1F, 0x4D, 0x78)}
    before = {1: 16, 2: 10}
    after = {1: 8, 2: 5}

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=before[level], after=after[level], line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, sizes[level], bold=True, color=colors[level])


def add_body(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, before=0, after=4, line=1.2)
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, size=10.5, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Tutor PPT 中文翻译与讲稿")
    set_run_font(run, 9, color=RGBColor(0x66, 0x66, 0x66))


def add_slide_block(document, slide_number, english_title, chinese_title, english_points, chinese_points, script):
    add_heading(document, f"第 {slide_number} 页：{english_title}", level=1)

    add_heading(document, "中文标题", level=2)
    add_body(document, chinese_title)

    add_heading(document, "英文页内要点", level=2)
    add_bullets(document, english_points)

    add_heading(document, "中文翻译", level=2)
    add_bullets(document, chinese_points)

    add_heading(document, "讲稿", level=2)
    add_body(document, script)


def add_summary_table(document):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(0.9)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(3.4)

    headers = ["页码", "英文标题", "本页作用"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color=RGBColor(0x1F, 0x4D, 0x78))

    rows = [
        ["1", "Text Diff in Dynamic History", "开场，定义这是一个工程化 diff 流水线。"],
        ["2", "Five processing layers", "说明主流程：预处理、拆块、对齐、细分、返回结果。"],
        ["3", "One project, several comparison modes", "讲不同内容类型采用不同 diff 粒度。"],
        ["4", "Supporting operations", "强调输入准备、清洗、安全与兜底。"],
        ["5", "Why this design is practical", "收束到可读性、稳定性、特化支持和产品化价值。"],
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def build_document():
    document = Document()
    configure_page(document)
    add_footer(document.sections[0])

    add_title(document, "Tutor 演示 PPT 中文翻译版与讲稿")
    add_subtitle(
        document,
        "对应文件：tutor-text-diff-presentation-en.pptx | 风格：全英文、朴素画面、适合 3 分钟左右讲解",
    )

    add_body(
        document,
        "这份 Word 文档包含两部分内容：第一，英文 PPT 的逐页中文翻译；第二，每一页对应的一段讲稿。这样可以直接拿来准备向 tutor 演示，也方便在正式讲之前快速对照英文页面与中文表达。"
    )

    add_heading(document, "整体结构", level=1)
    add_summary_table(document)

    add_slide_block(
        document,
        1,
        "Text Diff in Dynamic History",
        "Dynamic History 中的文本差异比较",
        [
            "This project does not rely on one diff function.",
            "It combines preprocessing, block alignment, type-specific comparison, and UI-friendly output.",
            "Input: Confluence storage HTML. Goal: readable and stable comparison. Output: HTML, blocks, and summary counters.",
        ],
        [
            "这个项目不是依赖一个单独的 diff 函数。",
            "它把预处理、块级对齐、按内容类型细分比较，以及适合前端消费的输出组合在一起。",
            "输入是 Confluence 的 storage HTML，目标是得到可读且稳定的比较结果，输出包括 HTML、结构化 blocks 和 summary 计数。",
        ],
        "这一页我先交代整体定位。这个项目的关键点不是写了一个神奇的 diff 函数，而是把多个步骤组合成一条完整的流水线。输入是 Confluence 的富文本存储格式，最终输出要既能让用户读懂，也要方便前端稳定渲染，所以它天然就是一个工程化方案。"
    )

    add_slide_block(
        document,
        2,
        "Five processing layers",
        "五层处理链",
        [
            "Preprocess rich content first.",
            "Split the page into comparable blocks.",
            "Use block-level LCS for alignment.",
            "Choose inline, line, or table-specific diff.",
            "Return HTML, blocks, summary, and limited flags.",
        ],
        [
            "先对富文本做预处理。",
            "再把页面拆成可比较的内容块。",
            "接着用块级 LCS 做对齐。",
            "然后根据类型选择行内 diff、按行 diff，或者表格专用 diff。",
            "最后返回 HTML、结构化 blocks、summary 和 limited 标记。",
        ],
        "第二页我会把主流程讲清楚。先处理 Confluence 自带的链接、宏和附件，再把页面切成段落、标题、列表、表格、代码块这些可比较单元。之后做块级对齐，找到哪些块是相同、删除、新增或修改。最后再根据块类型进入不同的比较逻辑，把结果统一返回给 UI。"
    )

    add_slide_block(
        document,
        3,
        "One project, several comparison modes",
        "一个项目，多种比较策略",
        [
            "Block-level LCS handles coarse alignment.",
            "Inline token diff handles normal text changes.",
            "Long-text fallback prevents the browser from freezing.",
            "Code blocks and tables use specialized logic.",
        ],
        [
            "块级 LCS 负责粗粒度对齐。",
            "行内 token diff 负责普通文本中的增删变化。",
            "长文本回退策略用于避免浏览器卡死。",
            "代码块和表格各自使用专门的比较逻辑。",
        ],
        "第三页的重点是说明，这里不是一种 diff 打天下。普通段落适合做 token 级别比较；内容太长时，先退回到句子或行的层级；代码块要保留缩进，所以按行比较更合理；表格如果结构没变，就做单元格高亮，如果结构变了，就并排展示。也就是说，算法的价值在于“选对粒度”。"
    )

    add_slide_block(
        document,
        4,
        "Supporting operations",
        "配套操作同样重要",
        [
            "src/index.js fetches page versions, attachment URLs, and author names.",
            "prepareConfluenceHtml normalizes storage-format HTML before comparison.",
            "Safe rendering filters risky tags and attributes.",
            "ComparisonPanel catches rendering failures and falls back safely.",
        ],
        [
            "src/index.js 会拉取页面版本、附件 URL 和作者信息。",
            "prepareConfluenceHtml 会在比较前规范化 storage HTML。",
            "安全渲染步骤会过滤风险标签与属性。",
            "ComparisonPanel 会接住渲染错误，并安全回退，而不是让页面崩溃。",
        ],
        "第四页我要强调，真正让这个功能可上线的，不只是 diff 算法本身。前面必须先拿到完整版本数据和附件地址，中间要把 storage HTML 转成安全、可渲染的内容，后面还要有错误兜底。也就是说，这个功能的稳定性来自算法和工程配套一起工作。"
    )

    add_slide_block(
        document,
        5,
        "Why this design is practical",
        "为什么这套设计是实用的",
        [
            "Readable for users.",
            "Stable for large content.",
            "Specialized for code and tables.",
            "Connected to real product behavior.",
        ],
        [
            "对用户来说更易读。",
            "面对大内容时更稳定。",
            "对代码块和表格有针对性的支持。",
            "和真实产品行为是连起来的，不只是实验性算法。",
        ],
        "最后一页我会做总结。这个设计的价值，不是它提出了新的 diff 理论，而是它把经典 LCS 思路真正落到了产品里。用户能看懂，浏览器跑得动，代码和表格不会被错误处理，而且整个前后端链路是闭合的。所以我会把它定义成一个 practical design，而不是只停留在算法层面的 demo。"
    )

    add_heading(document, "讲解节奏建议", level=1)
    add_bullets(
        document,
        [
            "第 1 页控制在 25 到 30 秒，快速定性项目。",
            "第 2 页和第 3 页是核心，合计大约 90 秒。",
            "第 4 页讲工程配套，约 30 到 40 秒。",
            "第 5 页用 20 到 30 秒收尾，强调 practical value。",
        ],
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
